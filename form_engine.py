# -*- coding: utf-8 -*-
from __future__ import annotations
from io import BytesIO
from pathlib import Path
from typing import Any
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from experiment_engine import result_summary
from report_rules import overall_conclusion, report_item

ROOT=Path(__file__).parent
TEMPLATE_DIR=ROOT/"templates"
SIGNATURE_DIR=ROOT/"data"/"signatures"
BLACK=RGBColor(0,0,0)


def _blacken(doc):
    for p in doc.paragraphs:
        for r in p.runs:r.font.color.rgb=BLACK
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:r.font.color.rgb=BLACK


def _save(doc):
    _blacken(doc);b=BytesIO();doc.save(b);b.seek(0);return b


def _setp(p,text,bold=False):
    p.clear();r=p.add_run(str(text));r.bold=bold;r.font.color.rgb=BLACK


def _prefix(p,prefix,value):
    if p.text.strip().startswith(prefix):_setp(p,prefix+("" if value is None else str(value)));return True
    return False


def _fill(table,data,header_rows=1):
    while len(table.rows)<header_rows+len(data):table.add_row()
    for i,vals in enumerate(data,start=header_rows):
        for j,v in enumerate(vals):
            if j<len(table.rows[i].cells):table.rows[i].cells[j].text="" if v is None else str(v)
    for i in range(header_rows+len(data),len(table.rows)):
        for cell in table.rows[i].cells:cell.text=""


def group_range(g):
    q=int(g.get("quantity") or 1)
    return f"{g['group_no']}-01" if q==1 else f"{g['group_no']}-01～{g['group_no']}-{q:02d}"


def commission_document(c,groups,tests,receiver_name):
    d=Document(TEMPLATE_DIR/"FORM_COMMISSION.docx")
    methods=json.loads(c.get("method_choices") or "[]")
    options=["YY/T 1936","YY 0300","YY 0621.1","YY 0621.2","YY/T 1702","GB 17168","GB/T 4340.1","GB/T 3851","GB/T 18876.1","YY/T 1937","YY 0270.1","T/GDMDMA 0003","YY 0710"]
    method_line1="  ".join(("☑" if x in methods else "□")+x for x in options[:7])
    method_line2="  ".join(("☑" if x in methods else "□")+x for x in options[7:])
    bad=[g for g in groups if g.get("condition")!="完好"]
    bad_note="；".join(f"{g['group_no']}:{g.get('condition_note','')}" for g in bad)
    for p in d.paragraphs:
        _prefix(p,"委托方名称：",c.get("client_name",""));_prefix(p,"委托方地址：",c.get("client_address",""))
        if p.text.strip().startswith("联系人："):_setp(p,f"联系人：{c.get('contact','')}    联系电话：{c.get('phone','')}    委托日期：{c.get('commission_date','')}")
        if "样品外观检查良好" in p.text:_setp(p,f"{'□' if bad else '☑'}样品外观检查良好； {'☑' if bad else '□'}样品外观异常。异常情况说明：{bad_note}")
        if p.text.strip().startswith("检测方法："):_setp(p,"检测方法：")
        elif "YY/T 1936" in p.text:_setp(p,method_line1)
        elif "GB/T 3851" in p.text:_setp(p,method_line2)
        if p.text.strip().startswith("1、标普检测资源不满足时"):
            yes=c.get("subcontract_allowed")=="是";_setp(p,f"1、标普检测资源不满足时，是否允许分包？ {'☑是  □否' if yes else '□是  ☑否'}")
        if p.text.strip().startswith("2、样品、相关资料是否保密"):_setp(p,"2、样品、相关资料是否保密？ □是  ☑无要求")
        if p.text.strip().startswith("报告载体："):_setp(p,f"报告载体：{c.get('report_medium','')}    符合性判定：{c.get('conformity_judgment','')}")
        if p.text.strip().startswith("考虑不确定度："):_setp(p,f"考虑不确定度：{c.get('uncertainty','')}    递送方式：{c.get('delivery_method','')}")
        if "CNAS" in p.text and "章" in p.text:_setp(p,f"加盖 CNAS 章：{c.get('cnas_mark','')}")
        if p.text.strip().startswith("检测能力："):_setp(p,f"检测能力：{c.get('capability','')}")
    tm={}
    for t in tests:tm.setdefault(t["group_no"],[]).append(t["experiment"])
    data=[]
    for i,g in enumerate(groups,1):
        prod=c.get("production_org_name","")+("（受委托生产企业）" if c.get("production_relation")=="受委托生产企业" else "")
        data.append([i,f"{g['sample_name']}（{g['model']}）",group_range(g),prod,"、".join(tm.get(g["group_no"],[])),g.get("quantity",1),g.get("notes","") or g.get("condition_note","")])
    _fill(d.tables[0],data)
    return _save(d)


def sample_register_document(c,groups,samples,tests,receiver_name):
    """Generate DLBP-CX-P10-R01 sample registration form.

    Column order follows the controlled template exactly:
    laboratory sample number, commissioning unit, sample name, model/specification,
    production unit, sample number/batch, inspection items, quantity, receiver,
    date and remarks.
    """
    d=Document(TEMPLATE_DIR/"FORM_SAMPLE_REGISTER.docx")
    gm={g["group_no"]:g for g in groups}
    tm={}
    for t in tests:
        tm.setdefault(t["group_no"],[]).append(t["experiment"])

    production_unit=c.get("production_org_name","")
    if c.get("production_relation")=="受委托生产企业" and production_unit:
        production_unit += "（受委托生产企业）"

    data=[]
    for s in samples:
        g=gm[s["group_no"]]
        remarks=s.get("condition_note","") or g.get("notes","") or ""
        data.append([
            s["sample_no"],
            c.get("client_name",""),
            s["sample_name"],
            s["model"],
            production_unit,
            g.get("product_no",""),
            "、".join(tm.get(s["group_no"],[])),
            1,
            receiver_name,
            c.get("commission_date",""),
            remarks,
        ])
    _fill(d.tables[0],data)
    return _save(d)


def loan_return_document(loans,user_names):
    d=Document(TEMPLATE_DIR/"FORM_SAMPLE_LOAN_RETURN.docx");data=[]
    for i,x in enumerate(loans,1):
        purpose=x.get("purpose") or "、".join(json.loads(x.get("experiments") or "[]"))
        data.append([i,x["sample_no"],user_names.get(x.get("borrower"),x.get("borrower","")),x.get("borrowed_at",""),purpose,x.get("returned_at",""),user_names.get(x.get("returned_by"),x.get("returned_by","")),x.get("return_note","") or x.get("issue_note","")])
    _fill(d.tables[0],data);return _save(d)


def _sig(meta):
    if not meta or not meta.get("image_file"):return None
    p=SIGNATURE_DIR/meta["image_file"];return p if p.exists() else None


def _set_existing_text(paragraph, text):
    """Replace text without removing the template paragraph/run elements."""
    runs=list(paragraph.runs)
    if not runs:
        runs=[paragraph.add_run("")]
    runs[0].text=str(text or "")
    runs[0].font.color.rgb=BLACK
    for run in runs[1:]:
        run.text=""
        run.font.color.rgb=BLACK


def _prefix_existing(paragraph, prefix, value):
    if paragraph.text.strip().startswith(prefix):
        _set_existing_text(paragraph,prefix+("" if value is None else str(value)))
        return True
    return False


def _set_cell_existing(cell, value):
    paragraphs=list(cell.paragraphs)
    paragraph=paragraphs[0] if paragraphs else cell.add_paragraph()
    _set_existing_text(paragraph,"" if value is None else str(value))
    for extra in paragraphs[1:]:
        _set_existing_text(extra,"")


def _fill_existing_rows(table, data, header_rows=1):
    """Fill only rows already present in the controlled report mother."""
    capacity=max(0,len(table.rows)-header_rows)
    for offset in range(capacity):
        values=data[offset] if offset<len(data) else []
        row=table.rows[header_rows+offset]
        for col,cell in enumerate(row.cells):
            _set_cell_existing(cell,values[col] if col<len(values) else "")


def _date_range(values):
    clean=sorted(dict.fromkeys(str(x)[:10] for x in values if x))
    if not clean:return ""
    return clean[0] if len(clean)==1 else f"{clean[0]}至{clean[-1]}"


def _range_text(values, suffix=""):
    nums=[]
    for value in values:
        try:nums.append(float(value))
        except Exception:pass
    if not nums:return ""
    lo,hi=min(nums),max(nums)
    def fmt(x):return f"{x:.3f}".rstrip("0").rstrip(".")
    return f"{fmt(lo)}{suffix}" if lo==hi else f"{fmt(lo)}～{fmt(hi)}{suffix}"


def report_document(c,groups,samples,tasks,records,report,user_names,signatures):
    d=Document(TEMPLATE_DIR/"FORM_REPORT.docx")
    names="、".join(dict.fromkeys(g["sample_name"] for g in groups));models="、".join(dict.fromkeys(g["model"] for g in groups));ranges="；".join(group_range(g) for g in groups);products="、".join(dict.fromkeys(g.get("product_no","") for g in groups if g.get("product_no")));prods=c.get("production_org_name","")+("（受委托生产企业）" if c.get("production_relation")=="受委托生产企业" else "");conds="；".join(f"{g['group_no']}:{g['condition']}" for g in groups)
    test_dates=[];report_items=[];equipment_map={};environment_by_location={}
    for t in tasks:
        rec=records.get(t["task_no"])
        if not rec:continue
        payload=rec.get("payload") or {}
        business=payload.get("business_record") or {}
        params=business.get("parameters") or payload.get("parameters") or {}
        rows=business.get("rows") or payload.get("data") or []
        common=payload.get("common") or {}
        snapshot=payload.get("configuration_snapshot") or {}
        kind=t.get("kind") or snapshot.get("kind") or "generic"
        item=report_item(kind,rows)
        item.update({
            "task":t,
            "standard":t.get("standard") or snapshot.get("standard") or t.get("method_code",""),
            "deviation":business.get("deviation") or payload.get("deviation") or "无",
        })
        report_items.append(item)
        test_dates.append(params.get("test_date") or common.get("test_date"))
        location=snapshot.get("default_location") or params.get("detection_location") or "未记录地点"
        env=environment_by_location.setdefault(location,{"temperature":[],"humidity":[],"other":[]})
        env["temperature"].extend([params.get("temperature_before"),params.get("temperature_after")])
        env["humidity"].extend([params.get("humidity_before"),params.get("humidity_after")])
        if params.get("environment_interference"):
            env["other"].append(str(params.get("environment_interference")))
        check_map={x.get("management_no"):x for x in business.get("equipment_checks") or []}
        for eq in snapshot.get("equipment") or []:
            no=eq.get("management_no","")
            value=dict(eq);value["usage_status"]=(check_map.get(no) or {}).get("status","正常")
            equipment_map.setdefault(no or eq.get("equipment_name",""),value)

    auto_statement="；".join(f"{g['group_no']}：接收状态{g.get('condition','完好')}，共{g.get('quantity',1)}件" for g in groups)
    auto_conclusion=overall_conclusion(report_items)
    for p in d.paragraphs:
        if p.text.strip().startswith("生 产 单 位："):
            _set_existing_text(p,f"生 产 单 位：{prods}    接 收 日 期：{c.get('commission_date','')}")
            continue
        _prefix_existing(p,"报告编号：",report.get("report_no",""));_prefix_existing(p,"委 托 单 位：",c.get("client_name",""));_prefix_existing(p,"地       址：",c.get("client_address",""));_prefix_existing(p,"样 品 名 称：",names);_prefix_existing(p,"型 号/规 格：",models);_prefix_existing(p,"样 品 编 号：",ranges);_prefix_existing(p,"产品编号/批号：",products);_prefix_existing(p,"生 产 单 位：",prods);_prefix_existing(p,"接 收 日 期：",c.get("commission_date",""));_prefix_existing(p,"接 收 状 态：",conds);_prefix_existing(p,"检 验 类 别：",report.get("report_category") or "委托检验");_prefix_existing(p,"报告发布日期：",report.get("publish_date") or "")
        _prefix_existing(p,"检验日期 ：",_date_range(test_dates))
        if p.text.strip().startswith("需说明的情况:"):_set_existing_text(p,"需说明的情况:"+(report.get("notes") or "无"))
        if p.text.strip().startswith("样品情况说明："):_set_existing_text(p,"样品情况说明："+(report.get("sample_statement") or auto_statement))
        if p.text.strip().startswith("检验结论："):_set_existing_text(p,"检验结论："+(report.get("conclusion") or auto_conclusion))

    equipment_rows=[]
    equipment=list(equipment_map.values())
    for item in equipment[:5]:
        equipment_rows.append([
            f"{item.get('equipment_name','')}（{item.get('management_no','')}）",
            item.get("model",""),
            item.get("measuring_range",""),
            item.get("calibration_certificate") or "台账未配置",
            item.get("traceability_agency") or "台账未配置",
            item.get("calibration_due") or (f"台账校准时间：{item.get('calibration_time')}" if item.get("calibration_time") else "台账未配置"),
        ])
    environment_rows=[
        [location,_range_text(data["temperature"]," ℃"),_range_text(data["humidity"]," %RH"),"、".join(dict.fromkeys(data["other"])) or "无"]
        for location,data in environment_by_location.items()
    ][:3]

    grouped={}
    for item in report_items:
        group_no=item["task"].get("group_no","")
        grouped.setdefault(group_no,[]).append(item)
    result_rows=[]
    for index,(group_no,items) in enumerate(grouped.items(),1):
        experiments="；".join(f"{x['task'].get('experiment','')}（{x['standard']}）" for x in items)
        requirements="；".join(dict.fromkeys(x["requirement"] for x in items))
        actual="；".join(x["result"] for x in items)
        conclusions=[x["conclusion"] for x in items]
        conclusion="不符合" if any(x in ("不符合","不合格") for x in conclusions) else ("符合" if all(x in ("符合","合格") for x in conclusions) else "仅描述结果")
        notes="；".join(dict.fromkeys(x["deviation"] for x in items if x["deviation"] not in ("","无"))) or "无"
        result_rows.append([index,f"{group_no} {experiments}",requirements,actual,conclusion,notes])
    if len(result_rows)>3:
        overflow=result_rows[2:]
        result_rows=result_rows[:2]+[[
            3,
            "；".join(str(x[1]) for x in overflow),
            "；".join(str(x[2]) for x in overflow),
            "；".join(str(x[3]) for x in overflow),
            "不符合" if any(x[4]=="不符合" for x in overflow) else "符合",
            "；".join(str(x[5]) for x in overflow),
        ]]
    if len(d.tables)>0:_fill_existing_rows(d.tables[0],equipment_rows)
    if len(d.tables)>1:_fill_existing_rows(d.tables[1],environment_rows)
    if len(d.tables)>2:
        standards="；".join(dict.fromkeys(item["standard"] for item in report_items if item["standard"]))
        _set_cell_existing(d.tables[2].rows[0].cells[2],standards)
        _fill_existing_rows(d.tables[2],result_rows,2)
    for label,field,signed in [("批 准 人","approver","approver_signed_at"),("核 验 员","verifier","verifier_signed_at"),("检 测 员","tester","tester_signed_at")]:
        u=report.get(field);name=user_names.get(u,u or "")
        for p in d.paragraphs:
            if p.text.strip().startswith(label):
                signed_date=str(report.get(signed) or "")[:10]
                _set_existing_text(p,f"{label}    {name}    {signed_date}")
                if report.get(signed):
                    path=_sig(signatures.get(u))
                    if path:
                        try:p.runs[0].add_picture(str(path),width=Inches(.85))
                        except:pass
                break
    return _save(d)
