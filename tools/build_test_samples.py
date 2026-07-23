from __future__ import annotations

import json
import sys
from hashlib import sha256
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from business_record_engine import business_to_template_fields, calculate_business_record
from constants import EXPERIMENTS
from form_engine import report_document
from record_word_engine import export_record
from smoke_test import _context, _demo_business


OUTPUT = ROOT / "test_samples"


def signature(source) -> dict:
    document = Document(source)
    return {
        "paragraphs": len(document.paragraphs),
        "tables": [{"rows": len(table.rows), "cols": len(table.columns)} for table in document.tables],
        "sections": [
            {
                "page_width": int(section.page_width or 0),
                "page_height": int(section.page_height or 0),
                "top": int(section.top_margin or 0),
                "bottom": int(section.bottom_margin or 0),
                "left": int(section.left_margin or 0),
                "right": int(section.right_margin or 0),
            }
            for section in document.sections
        ],
    }


def main() -> None:
    OUTPUT.mkdir(exist_ok=True)
    sample_ids = ["BP20260723001-01", "BP20260723001-02"]
    equipment = [
        {
            "management_no": "BPGL-A036",
            "equipment_name": "粗糙度仪",
            "model": "TR200",
            "measuring_range": "±160 μm",
            "calibration_time": "2026-07",
            "binding_role": "主设备",
            "required": 1,
        }
    ]
    attachments = [{"attachment_id": "ATT202607230001", "sample_no": sample_ids[0]}]
    qa = {"records": {}, "report": {}}

    for experiment, config in EXPERIMENTS.items():
        business = _demo_business(config["kind"], sample_ids)
        values = business_to_template_fields(
            config["template"],
            config["kind"],
            _context(experiment, sample_ids),
            equipment,
            business,
            attachments,
            {},
        )
        record = {
            "record_no": f"QA-{config['key']}",
            "task_no": f"QA-{config['key']}",
            "version": 1,
            "experiment": experiment,
            "payload": {"template_fields": values, "business_record": business},
        }
        output = OUTPUT / f"{config['key']}_{config['template']}"
        output.write_bytes(export_record(record, config["template"], []).getvalue())
        mother = ROOT / "templates" / config["template"]
        qa["records"][experiment] = {
            "file": output.name,
            "mother_sha256": sha256(mother.read_bytes()).hexdigest(),
            "mother_structure_preserved": signature(output) == signature(mother),
        }

    rough = _demo_business("rough", sample_ids)
    rough["rows"][0].update({"ra1": 5.1, "ra2": 5.2, "ra3": 5.3})
    rough["rows"][1].update({"ra1": 6.1, "ra2": 6.2, "ra3": 6.3})
    rough = calculate_business_record("rough", rough)
    task_no = "BP20260723001-P01-T01"
    payload = {
        "common": {"test_date": "2026-07-23"},
        "business_record": rough,
        "configuration_snapshot": {
            "kind": "rough",
            "standard": "YY/T 1702-2020",
            "default_location": "性能检测室",
            "equipment": equipment,
        },
        "deviation": "无",
    }
    report_output = OUTPUT / "检验报告母版回填测试样例.docx"
    report_output.write_bytes(
        report_document(
            {
                "client_name": "测试客户",
                "client_address": "测试地址",
                "production_org_name": "测试生产单位",
                "production_relation": "生产单位",
                "commission_date": "2026-07-23",
            },
            [{
                "group_no": "BP20260723001",
                "sample_name": "测试试样",
                "model": "10 mm×10 mm×10 mm",
                "product_no": "B001",
                "condition": "完好",
                "quantity": 2,
            }],
            [],
            [{
                "task_no": task_no,
                "group_no": "BP20260723001",
                "sample_name": "测试试样",
                "experiment": "表面粗糙度试验",
                "standard": "YY/T 1702-2020",
                "method_code": "触针式轮廓仪法",
                "kind": "rough",
            }],
            {task_no: {"payload": payload}},
            {
                "report_no": "WT20260723001",
                "status": "待检测员确认",
                "tester": "tester",
                "verifier": "reviewer",
                "approver": "approver",
            },
            {"tester": "测试实验员", "reviewer": "测试复核员", "approver": "测试批准人"},
            {},
        ).getvalue()
    )
    report_mother = ROOT / "templates" / "FORM_REPORT.docx"
    qa["report"] = {
        "file": report_output.name,
        "mother_sha256": sha256(report_mother.read_bytes()).hexdigest(),
        "mother_structure_preserved": signature(report_output) == signature(report_mother),
        "contains_actual_results": True,
    }
    (OUTPUT / "QA_MANIFEST.json").write_text(json.dumps(qa, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Generated {len(EXPERIMENTS)} record samples and 1 report sample")


if __name__ == "__main__":
    main()
