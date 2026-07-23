from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DIR = ROOT / "templates"


def unique_cells(row):
    seen = set()
    output = []
    for index, cell in enumerate(row.cells):
        marker = id(cell._tc)
        if marker in seen:
            continue
        seen.add(marker)
        output.append((index, " ".join(cell.text.split())))
    return output


def audit_document(path: Path) -> dict:
    document = Document(path)
    sections = []
    for section in document.sections:
        sections.append(
            {
                "page_width": int(section.page_width or 0),
                "page_height": int(section.page_height or 0),
                "top_margin": int(section.top_margin or 0),
                "bottom_margin": int(section.bottom_margin or 0),
                "left_margin": int(section.left_margin or 0),
                "right_margin": int(section.right_margin or 0),
                "header_distance": int(section.header_distance or 0),
                "footer_distance": int(section.footer_distance or 0),
                "header": [" ".join(p.text.split()) for p in section.header.paragraphs],
                "footer": [" ".join(p.text.split()) for p in section.footer.paragraphs],
            }
        )
    return {
        "file": path.name,
        "paragraphs": [
            {"index": index, "text": " ".join(paragraph.text.split())}
            for index, paragraph in enumerate(document.paragraphs)
            if paragraph.text.strip()
        ],
        "tables": [
            {
                "index": table_index,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "cells": [
                    {
                        "row": row_index,
                        "values": unique_cells(row),
                    }
                    for row_index, row in enumerate(table.rows)
                ],
            }
            for table_index, table in enumerate(document.tables)
        ],
        "sections": sections,
    }


def main() -> None:
    destination = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("/tmp/bplab_template_audit.json")
    files = sorted(TEMPLATE_DIR.glob("RECORD_*.docx")) + [TEMPLATE_DIR / "FORM_REPORT.docx"]
    destination.write_text(
        json.dumps([audit_document(path) for path in files], ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(destination)


if __name__ == "__main__":
    main()
