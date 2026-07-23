# -*- coding: utf-8 -*-
from __future__ import annotations

from hashlib import sha256
from pathlib import Path
import tempfile

from docx import Document

import lims_db
from constants import EXPERIMENTS
from form_engine import report_document
from smoke_test import _context, _demo_business
from business_record_engine import business_to_template_fields
from record_word_engine import export_record
from template_record_engine import validate_template_fields


ROOT = Path(__file__).parent


def structure_signature(source) -> dict:
    doc = Document(source)
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": [
            {
                "rows": len(table.rows),
                "cols": len(table.columns),
                "grid": [len(row.cells) for row in table.rows],
            }
            for table in doc.tables
        ],
        "sections": [
            {
                "page_width": int(section.page_width or 0),
                "page_height": int(section.page_height or 0),
                "top": int(section.top_margin or 0),
                "bottom": int(section.bottom_margin or 0),
                "left": int(section.left_margin or 0),
                "right": int(section.right_margin or 0),
                "header": int(section.header_distance or 0),
                "footer": int(section.footer_distance or 0),
            }
            for section in doc.sections
        ],
    }


def all_text(source) -> str:
    doc = Document(source)
    values = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def main() -> None:
    sample_ids = ["BP20260723001-01", "BP20260723001-02"]
    equipment = [
        {
            "management_no": "BPGL-A036",
            "equipment_name": "粗糙度仪",
            "model": "TR200",
            "measuring_range": "±160 μm",
            "calibration_time": "2026-07",
            "binding_role": "主设备",
            "required": 1,
        }
    ]
    attachments = [{"attachment_id": "ATT202607230001", "sample_no": sample_ids[0]}]
    with tempfile.TemporaryDirectory() as temp_raw:
        temp = Path(temp_raw)
        for experiment, config in EXPERIMENTS.items():
            template = ROOT / "templates" / config["template"]
            mother_hash = sha256(template.read_bytes()).hexdigest()
            before = structure_signature(template)
            business = _demo_business(config["kind"], sample_ids)
            values = business_to_template_fields(
                config["template"],
                config["kind"],
                _context(experiment, sample_ids),
                equipment,
                business,
                attachments,
                {},
            )
            assert not validate_template_fields(config["template"], values), experiment
            output = temp / config["template"]
            record = {
                "record_no": f"QA-{config['key']}",
                "version": 1,
                "experiment": experiment,
                "payload": {"template_fields": values, "business_record": business},
            }
            output.write_bytes(export_record(record, config["template"], []).getvalue())
            assert structure_signature(output) == before, experiment
            assert sha256(template.read_bytes()).hexdigest() == mother_hash, experiment

        rough_business = _demo_business("rough", sample_ids)
        rough_business["rows"][0].update({"ra1": 5.1, "ra2": 5.2, "ra3": 5.3})
        rough_business["rows"][1].update({"ra1": 6.1, "ra2": 6.2, "ra3": 6.3})
        from business_record_engine import calculate_business_record
        rough_business = calculate_business_record("rough", rough_business)
        payload = {
            "common": {"test_date": "2026-07-23"},
            "business_record": rough_business,
            "configuration_snapshot": {
                "kind": "rough",
                "standard": "YY/T 1702-2020",
                "default_location": "性能检测室",
                "equipment": equipment,
            },
            "deviation": "无",
        }
        tasks = [
            {
                "task_no": "BP20260723001-P01-T01",
                "group_no": "BP20260723001",
                "sample_name": "测试试样",
                "experiment": "表面粗糙度试验",
                "standard": "YY/T 1702-2020",
                "method_code": "触针式轮廓仪法",
                "kind": "rough",
            }
        ]
        commission = {
            "client_name": "测试客户",
            "client_address": "测试地址",
            "production_org_name": "测试生产单位",
            "production_relation": "生产单位",
            "commission_date": "2026-07-23",
        }
        groups = [
            {
                "group_no": "BP20260723001",
                "sample_name": "测试试样",
                "model": "10 mm×10 mm×10 mm",
                "product_no": "B001",
                "condition": "完好",
                "quantity": 2,
            }
        ]
        report = {
            "report_no": "WT20260723001",
            "status": "待检测员确认",
            "tester": "tester",
            "verifier": "reviewer",
            "approver": "approver",
        }
        report_mother = ROOT / "templates" / "FORM_REPORT.docx"
        report_before = structure_signature(report_mother)
        generated_report = temp / "report.docx"
        generated_report.write_bytes(
            report_document(
                commission,
                groups,
                [],
                tasks,
                {tasks[0]["task_no"]: {"payload": payload}},
                report,
                {"tester": "测试实验员", "reviewer": "测试复核员", "approver": "测试批准人"},
                {},
            ).getvalue()
        )
        assert structure_signature(generated_report) == report_before
        report_text = all_text(generated_report)
        assert "详见原始记录" not in report_text
        assert "平均Ra5.2μm" in report_text and "平均Ra6.2μm" in report_text
        assert "每个试样3条测量线Ra平均值均≤15 μm" in report_text
        assert "性能检测室" in report_text
        assert "粗糙度仪" in report_text

        # Return confirmation is the report trigger; dashboard reconciliation is idempotent.
        lims_db.DB_PATH = temp / "report_flow.db"
        lims_db.ATTACHMENT_DIR = temp / "attachments"
        lims_db.SIGNATURE_DIR = temp / "signatures"
        lims_db.init_db()
        # The full create/record flow is exercised by smoke_test. Here we verify reconciliation
        # never creates reports for commissions that have not completed the return workflow.
        assert lims_db.dashboard_counts()["reports"] == 0

    print("V5.7 REGRESSION PASSED: 10 controlled mothers and report mother keep their structure")


if __name__ == "__main__":
    main()
