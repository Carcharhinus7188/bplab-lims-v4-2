# -*- coding: utf-8 -*-
from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import RGBColor

from template_record_engine import TEMPLATE_DIR, fill_exact_template

BLACK = RGBColor(0, 0, 0)


def _fallback(record):
    """Fallback is used only for a newly added experiment without a controlled template."""
    payload = record.get("payload", {})
    doc = Document()
    doc.add_heading(record.get("experiment") or "实验原始记录", 0)
    doc.add_paragraph("当前实验尚未配置受控原始记录模板。正式提交前应由管理员上传并发布模板版本。")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "字段"
    table.rows[0].cells[1].text = "记录值"
    for key, value in (payload.get("template_fields") or {}).items():
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = str(value)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = BLACK
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def export_record(record, template_name, changes):
    """Fill the controlled DOCX directly without adding, deleting or rebuilding its layout.

    The exported document contains only the original template content and values written into
    the template's existing cells. Equipment details and attachment indexes are not appended.
    """
    if not template_name or not (TEMPLATE_DIR / template_name).exists():
        return _fallback(record)

    changed_keys = set()
    if int(record.get("version", 1) or 1) > 1:
        for item in changes or []:
            field_name = str(item.get("field_name", ""))
            if item.get("action") == "字段修改" and field_name.startswith("template_fields."):
                changed_keys.add(field_name.split("template_fields.", 1)[1])

    values = (record.get("payload") or {}).get("template_fields") or {}
    doc = fill_exact_template(template_name, values, changed_keys)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
