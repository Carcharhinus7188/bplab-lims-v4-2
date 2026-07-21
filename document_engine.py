# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document
from docx.shared import Inches,Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io,json
ROOT=Path(__file__).parent;T=ROOT/"templates";SIG=ROOT/"data"/"signatures"
def setp(p,label,value):
 if label in p.text:
  p.text=f"{label}{value}"
def fill_table(table,rows,start=1):
 for i,row in enumerate(rows,start):
  if i>=len(table.rows):table.add_row()
  for j,v in enumerate(row):
   if j<len(table.rows[i].cells):table.rows[i].cells[j].text=str(v)
def commission_doc(c,samples,tasks):
 d=Document(T/"FORM_COMMISSION.docx")
 for p in d.paragraphs:
  setp(p,"委托方名称：",c["customer_name"]);setp(p,"委托方地址：",c["address"])
  if p.text.startswith("联系人："):p.text=f"联系人：{c['contact']}    联系电话：{c['phone']}    委托日期：{c['commission_date']}"
  if "样品外观检查良好" in p.text:p.text=("☑样品外观检查良好； □样品外观异常。" if c["condition"]=="完好" else f"□样品外观检查良好； ☑样品外观异常。异常情况说明：{c['condition_note']}")
  if p.text.startswith("1、标普检测资源"):p.text=f"1、标普检测资源不满足时，是否允许分包？ {'☑是 □否' if c['subcontract']=='是' else '□是 ☑否'}"
  if p.text.startswith("2、样品、相关资料"):p.text=f"2、样品、相关资料是否保密？ {'☑是 □无要求' if c['confidential']=='是' else '□是 ☑无要求'}"
 methods=sorted(set(t["standard"] for t in tasks if t["standard"]))
 for p in d.paragraphs:
  if p.text.startswith("检测方法："):p.text="检测方法："+ "；".join(methods)
 table=d.tables[0];by={}
 for t in tasks:by.setdefault(t["sample_no"],[]).append(t["experiment"])
 fill_table(table,[[i,s["sample_name"],s["sample_no"],s["production_unit"],"、".join(by.get(s["sample_no"],[])),1,s["condition_note"]] for i,s in enumerate(samples,1)])
 return save(d)
def register_doc(c,samples,tasks):
 d=Document(T/"FORM_SAMPLE_REGISTER.docx");by={}
 for t in tasks:by.setdefault(t["sample_no"],[]).append(t["experiment"])
 fill_table(d.tables[0],[[s["sample_no"],c["customer_name"],s["sample_name"],s["model"],s["product_no"],"、".join(by.get(s["sample_no"],[])),1,c["created_by"],c["commission_date"],s["condition_note"]] for s in samples])
 return save(d)
def loan_doc(loans):
 d=Document(T/"FORM_SAMPLE_LOAN_RETURN.docx")
 fill_table(d.tables[0],[[i,x["sample_no"],x["borrower"],x["borrowed_at"],x["experiment"],x.get("returned_at") or "",x.get("returned_by") or "",x.get("return_note") or x.get("issue_note") or ""] for i,x in enumerate(loans,1)])
 return save(d)
def result_summary(payload):
 data=payload.get("data",[]);out=[]
 for r in data[:6]:
  vals=[f"{k}:{v}" for k,v in r.items() if k not in("试样编号","样品编号") and v not in("",None)]
  out.append("，".join(vals[-3:]))
 return "；".join(out)
def report_doc(c,samples,tasks,report,records,users,signatures):
 d=Document(T/"FORM_REPORT.docx");s0=samples[0] if samples else {}
 mapping={"报告编号：":report["report_no"],"委 托 单 位：":c["customer_name"],"地       址：":c["address"],"样 品 名 称：":"、".join(sorted(set(x["sample_name"] for x in samples))),"型 号/规 格：":"、".join(sorted(set(x["model"] for x in samples))),"样 品 编 号：":"、".join(x["sample_no"] for x in samples),"产品编号/批号：":"、".join(sorted(set(x["product_no"] for x in samples if x["product_no"]))),"接 收 状 态：":c["condition"],"检 验 类 别：":"委托检验","报告发布日期：":report.get("publish_date") or "","检验日期 ：":""}
 for p in d.paragraphs:
  for k,v in mapping.items():setp(p,k,v)
  if p.text.startswith("生 产 单 位："):p.text=f"生 产 单 位：{s0.get('production_unit','')}    接 收 日 期：{c['commission_date']}"
  if p.text.startswith("样品情况说明："):p.text="样品情况说明："+(report.get("sample_statement") or "")
  if p.text.startswith("检验结论："):p.text="检验结论："+(report.get("conclusion") or "")
 # equipment/environment/results
 eq=[];env=[];results=[]
 for t in tasks:
  r=records.get(t["task_no"])
  if not r:continue
  p=r["payload"];cm=p.get("common",{})
  eq.append([cm.get("equipment",""),"",cm.get("calibration",""),"","",""])
  env.append([cm.get("location",""),cm.get("temperature",""),cm.get("humidity",""),""])
  results.append([len(results)+1,t["experiment"],t["standard"],result_summary(p),next((x.get("判定","") for x in p.get("data",[]) if x.get("判定")), ""), ""])
 if d.tables:fill_table(d.tables[0],eq)
 if len(d.tables)>1:fill_table(d.tables[1],env)
 if len(d.tables)>2:fill_table(d.tables[2],results,start=2)
 # signatures: add to label paragraphs
 stages=[("批 准 人",report.get("approver")),("核 验 员",report.get("verifier")),("检 测 员",report.get("tester"))]
 for label,u in stages:
  info=users.get(u,{})
  for p in d.paragraphs:
   if label in p.text:
    p.add_run("  "+info.get("display_name",u or ""))
    img=signatures.get(u)
    if img and Path(img).exists():
     try:p.add_run().add_picture(str(img),width=Inches(0.9))
     except:pass
 return save(d)
def save(d):
 b=io.BytesIO();d.save(b);b.seek(0);return b
