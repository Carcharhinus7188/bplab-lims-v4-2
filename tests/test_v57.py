from __future__ import annotations

import os
import sqlite3
import tempfile
import unittest
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

from bplab.config import EXPERIMENTS
from bplab.db import execute, init_db, json_load, query, query_one
from bplab.demo import seed_full_demo
from bplab.excel_trace import export_trace_excel
from bplab.record_export import _structure_signature, export_record
from bplab.report_export import export_report
from bplab.services import auto_build_missing_reports, report_release_gaps, task_detail


class V57FullWorkflowTest(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.temp = tempfile.TemporaryDirectory(prefix="bplab-v57-")
        cls.root = Path(cls.temp.name)
        cls.db = cls.root / "test.db"
        os.environ["BPLAB_DB_PATH"] = str(cls.db)
        init_db(cls.db)
        cls.commission_id = seed_full_demo(cls.db)
        cls.output = cls.root / "output"
        cls.output.mkdir()

    @classmethod
    def tearDownClass(cls):
        cls.temp.cleanup()

    def test_all_ten_experiments_complete_and_locked(self):
        tasks = query(
            "SELECT experiment_code,status,returned_at,judgment FROM tasks WHERE commission_id=?",
            (self.commission_id,),
            self.db,
        )
        self.assertEqual(set(EXPERIMENTS), {row["experiment_code"] for row in tasks})
        self.assertEqual(10, len(tasks))
        self.assertTrue(all(row["status"] == "locked" for row in tasks))
        self.assertTrue(all(row["returned_at"] for row in tasks))
        self.assertTrue(all(row["judgment"] for row in tasks))

    def test_report_is_auto_built_and_dashboard_bug_is_fixed(self):
        report = query_one("SELECT * FROM reports WHERE commission_id=?", (self.commission_id,), self.db)
        self.assertIsNotNone(report)
        self.assertEqual("待检测员确认", report["status"])
        execute("DELETE FROM reports WHERE id=?", (report["id"],), self.db)
        created = auto_build_missing_reports(self.db)
        self.assertEqual(1, len(created))
        rebuilt = query_one("SELECT * FROM reports WHERE commission_id=?", (self.commission_id,), self.db)
        self.assertIsNotNone(rebuilt)

    def test_report_contains_real_results_and_no_see_original_record(self):
        report = query_one("SELECT * FROM reports WHERE commission_id=?", (self.commission_id,), self.db)
        report["snapshot"] = json_load(report["snapshot_json"], {})
        for task in report["snapshot"]["tasks"]:
            result = task["calculations"]["report_result"]
            self.assertTrue(result)
            self.assertNotIn("详见原始记录", result)
            self.assertTrue(task["calculations"]["standard_requirement"])
        path, manifest = export_report(report, official=False, output_dir=self.output)
        self.assertTrue(path.exists())
        self.assertTrue(manifest["structure_preserved"])
        self.assertFalse(manifest["contains_see_original_record"])
        full_text = "\n".join(
            cell.text
            for table in Document(path).tables
            for row in table.rows
            for cell in row.cells
        )
        self.assertNotIn("详见原始记录", full_text)
        self.assertIn("表面粗糙度试验", full_text)
        self.assertIn("维氏硬度试验", full_text)
        self.assertIn("牙科材料色稳定性试验", full_text)

    def test_controlled_record_templates_preserve_structure(self):
        controlled = 0
        for row in query("SELECT id FROM tasks WHERE commission_id=? ORDER BY id", (self.commission_id,), self.db):
            task = task_detail(row["id"], self.db)
            if task["config"]["template_status"] != "controlled":
                continue
            controlled += 1
            template = Path(__file__).resolve().parents[1] / "templates" / "records" / task["config"]["record_template"]
            before = _structure_signature(Document(template))
            output, manifest = export_record(task, self.output)
            after = _structure_signature(Document(output))
            self.assertEqual(before, after, task["experiment_code"])
            self.assertTrue(manifest["structure_preserved"])
        self.assertEqual(7, controlled)

    def test_original_measurements_are_written_not_only_averages(self):
        bending = query_one("SELECT id FROM tasks WHERE experiment_code='bending'", path=self.db)
        task = task_detail(bending["id"], self.db)
        output, _ = export_record(task, self.output)
        doc = Document(output)
        result_table = doc.tables[4]
        self.assertIn("320", result_table.cell(1, 7).text)
        self.assertIn("850", result_table.cell(1, 8).text)
        vickers = query_one("SELECT id FROM tasks WHERE experiment_code='vickers'", path=self.db)
        output, _ = export_record(task_detail(vickers["id"], self.db), self.output)
        doc = Document(output)
        self.assertIn("460", doc.tables[4].cell(1, 2).text)
        self.assertIn("461", doc.tables[4].cell(1, 3).text)

    def test_excel_trace_has_required_sheets(self):
        path = export_trace_excel(self.commission_id, self.output, self.db)
        wb = load_workbook(path, read_only=True)
        required = {"使用说明", "实验总台账", "附件索引", "图片粘贴区", "修改记录", "复核记录"}
        self.assertTrue(required.issubset(set(wb.sheetnames)))
        self.assertEqual(16, len(wb.sheetnames))

    def test_attachment_schema_has_no_software_or_equipment_columns(self):
        with sqlite3.connect(self.db) as conn:
            columns = {row[1] for row in conn.execute("PRAGMA table_info(attachments)").fetchall()}
        self.assertNotIn("software", columns)
        self.assertNotIn("equipment", columns)
        self.assertIn("sha256", columns)

    def test_equipment_change_does_not_rewrite_history_snapshot(self):
        task_row = query_one("SELECT id FROM tasks WHERE experiment_code='vickers'", path=self.db)
        before = task_detail(task_row["id"], self.db)["equipment"]
        execute("UPDATE equipment SET name='新设备名称',version=version+1 WHERE management_no='BPGL-A035'", path=self.db)
        after = task_detail(task_row["id"], self.db)["equipment"]
        self.assertEqual(before, after)
        self.assertTrue(any(item["name"] == "数显维氏硬度计" for item in after))

    def test_release_gaps_are_explicit_not_silently_blank(self):
        report = query_one("SELECT id FROM reports WHERE commission_id=?", (self.commission_id,), self.db)
        gaps = report_release_gaps(report["id"], self.db)
        self.assertTrue(gaps)
        self.assertTrue(any("证书编号" in gap for gap in gaps))


if __name__ == "__main__":
    unittest.main(verbosity=2)

