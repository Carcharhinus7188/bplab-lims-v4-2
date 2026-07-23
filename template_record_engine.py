# -*- coding: utf-8 -*-
from __future__ import annotations

from copy import deepcopy
from functools import lru_cache
from pathlib import Path
import difflib
import re
from typing import Any

from docx import Document
from docx.document import Document as _Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).parent
TEMPLATE_DIR = ROOT / "templates"
BLACK = RGBColor(0, 0, 0)
RED = RGBColor(255, 0, 0)

BLANK_RE = re.compile(r"_{2,}|＿{2,}|…{2,}")
SPACE_RE = re.compile(r"\s+")
DATE_RE = re.compile(r"(\d{4})[-/.年](\d{1,2})[-/.月](\d{1,2})")
TIME_RE = re.compile(r"(\d{1,2}):(\d{1,2})(?::(\d{1,2}))?")


def _clean(value: Any) -> str:
    return SPACE_RE.sub(" ", str(value or "").replace("\xa0", " ")).strip()


def _norm(value: Any) -> str:
    return re.sub(r"[\s：:：/\\、，,；;（）()\[\]【】\-—_]+", "", str(value or "")).lower()


def _contains_marker(text: str) -> bool:
    value = _clean(text)
    return not value or bool(BLANK_RE.search(value)) or "□" in value or "☐" in value


def _body_table_sections(doc: _Document) -> list[str]:
    sections: list[str] = []
    last_text = ""
    table_index = 0
    for child in doc._element.body.iterchildren():
        if child.tag == qn("w:p"):
            text = _clean(Paragraph(child, doc._body).text)
            if text:
                last_text = text
        elif child.tag == qn("w:tbl"):
            sections.append(last_text or f"表{table_index + 1}")
            table_index += 1
    return sections


def _unique_row_cells(row) -> list[tuple[int, Any, str]]:
    result: list[tuple[int, Any, str]] = []
    seen = set()
    for col_index, cell in enumerate(row.cells):
        if cell._tc in seen:
            continue
        seen.add(cell._tc)
        result.append((col_index, cell, _clean(cell.text)))
    return result


def _table_headers(table) -> dict[int, str]:
    if not table.rows:
        return {}
    return {col_index: text for col_index, _, text in _unique_row_cells(table.rows[0])}


def _field_label(
    table_index: int,
    row_index: int,
    col_index: int,
    row_cells: list[tuple[int, Any, str]],
    headers: dict[int, str],
    template_text: str,
) -> tuple[str, str, str]:
    row_label = ""
    immediate_left = ""
    for candidate_col, _, candidate_text in row_cells:
        if candidate_col >= col_index:
            break
        if candidate_text:
            row_label = candidate_text
            immediate_left = candidate_text
    col_header = headers.get(col_index, "")

    meaningful_template = re.sub(r"[_＿]+", "", template_text).strip(" /：:；;")
    prefix = ""
    if template_text and meaningful_template and not template_text.startswith("□"):
        prefix = re.split(r"_{2,}|＿{2,}|□", template_text, maxsplit=1)[0].strip(" /：:；;")

    if not template_text and immediate_left and not _contains_marker(immediate_left):
        label = immediate_left
    elif col_header and not _contains_marker(col_header):
        label = col_header
    elif prefix:
        label = prefix
    elif row_label and not _contains_marker(row_label):
        label = row_label
    else:
        label = f"表{table_index + 1}第{row_index + 1}行第{col_index + 1}列"

    if row_index > 0 and col_header and label == col_header:
        label = f"{label}（第{row_index}条）"
    return label, row_label, col_header


def _infer_input_type(original: str) -> str:
    if "□" in original or "☐" in original:
        return "checkbox"
    if "年" in original and "月" in original and "日" in original and BLANK_RE.search(original):
        return "date"
    if ":" in original and BLANK_RE.search(original):
        return "time"
    if len(original) > 45:
        return "textarea"
    return "text"


@lru_cache(maxsize=128)
def template_manifest(template_name: str) -> list[dict[str, Any]]:
    path = TEMPLATE_DIR / template_name
    if not path.exists():
        return []
    doc = Document(path)
    sections = _body_table_sections(doc)
    fields: list[dict[str, Any]] = []
    for table_index, table in enumerate(doc.tables):
        headers = _table_headers(table)
        seen_cells = set()
        for row_index, row in enumerate(table.rows):
            row_cells = _unique_row_cells(row)
            for col_index, cell, template_text in row_cells:
                if cell._tc in seen_cells:
                    continue
                seen_cells.add(cell._tc)
                if row_index == 0 and not _contains_marker(template_text):
                    continue
                if not _contains_marker(template_text):
                    continue
                if row_index == 0 and not template_text:
                    left_text = next((text for c, _, text in reversed(row_cells) if c < col_index and text), "")
                    if not left_text:
                        continue
                label, row_label, col_header = _field_label(
                    table_index, row_index, col_index, row_cells, headers, template_text
                )
                fields.append({
                    "key": f"t{table_index}_r{row_index}_c{col_index}",
                    "section": sections[table_index] if table_index < len(sections) else f"表{table_index + 1}",
                    "table": table_index,
                    "row": row_index,
                    "col": col_index,
                    "label": label,
                    "row_label": row_label,
                    "col_header": col_header,
                    "template_text": template_text,
                    "input_type": _infer_input_type(template_text),
                    "position": f"表{table_index + 1}-R{row_index + 1}C{col_index + 1}",
                })
    return fields


def _select_checkbox(text: str, preferred: str) -> str:
    value = str(text or "")
    if not preferred or "□" not in value:
        return value
    # Handle optional spaces between the box and option text.
    pattern = re.compile(r"□\s*" + re.escape(preferred))
    if pattern.search(value):
        return pattern.sub(lambda m: "☑" + m.group(0)[1:], value, count=1)
    return value


def _select_other(text: str, replacement: str) -> str:
    value = str(text or "")
    if "其他" not in value:
        return value
    value = re.sub(r"□\s*其他", lambda m: "☑" + m.group(0)[1:], value, count=1)
    return BLANK_RE.sub(replacement, value, count=1)


def _fill_date(original: str, raw: str) -> str:
    match = DATE_RE.search(str(raw or ""))
    if not match:
        return original
    year, month, day = match.groups()
    groups = list(BLANK_RE.finditer(original))
    if len(groups) < 3:
        return BLANK_RE.sub(str(raw), original, count=1)
    replacements = [year, f"{int(month):02d}", f"{int(day):02d}"]
    output = original
    offset = 0
    for marker, replacement in zip(groups[:3], replacements):
        start, end = marker.start() + offset, marker.end() + offset
        output = output[:start] + replacement + output[end:]
        offset += len(replacement) - (marker.end() - marker.start())
    return output


def _fill_time(original: str, raw: str) -> str:
    match = TIME_RE.search(str(raw or ""))
    if not match:
        return original
    hour, minute, second = match.groups()
    replacements = [f"{int(hour):02d}", f"{int(minute):02d}"]
    if second is not None:
        replacements.append(f"{int(second):02d}")
    output = original
    groups = list(BLANK_RE.finditer(original))
    offset = 0
    for marker, replacement in zip(groups, replacements):
        start, end = marker.start() + offset, marker.end() + offset
        output = output[:start] + replacement + output[end:]
        offset += len(replacement) - (marker.end() - marker.start())
    return output


def _compose_cell_text(original: str, raw_value: Any) -> str:
    raw = str(raw_value or "").strip()
    if not original:
        return raw
    if not raw:
        return original
    # A complete edited checkbox string is stored as-is.
    if ("□" in original or "☐" in original) and ("☑" in raw or "□" in raw):
        return raw
    if "年" in original and "月" in original and "日" in original and BLANK_RE.search(original):
        return _fill_date(original, raw)
    if ":" in original and BLANK_RE.search(original) and TIME_RE.search(raw):
        return _fill_time(original, raw)
    if BLANK_RE.search(original):
        return BLANK_RE.sub(raw, original, count=1)
    return raw


EQUIPMENT_SYNONYMS: dict[str, list[str]] = {
    "万能试验机": ["电子万能试验机"],
    "2000n力传感器": ["电子万能试验机"],
    "烘箱": ["电热恒温干燥箱"],
    "温度计": ["高精度铂电阻温度检测仪", "温湿度计"],
    "测量平台升降支架": ["大理石平台"],
    "d65标准光源灯箱": ["D65对色灯箱"],
    "x射线机": ["医用X射线限束器"],
    "打印设备": ["干式激光成像仪"],
    "标准密度片": ["密度片"],
    "标准硬度块": ["标准维氏硬度块"],
    "数显维氏硬度计": ["数显维氏硬度计"],
    "影像测量仪": ["影像测量仪"],
    "粗糙度仪": ["粗糙度仪"],
    "粗糙度标准块": ["粗糙度标准块"],
    "数显游标卡尺": ["数显游标卡尺"],
    "挠度计": ["挠度计"],
    "热膨胀测试仪": ["热膨胀测试仪"],
    "照度计": ["照度计"],
    "秒表": ["电子秒表"],
    "放大镜": ["放大镜"],
    "耐光色稳定性测试仪": ["耐光色稳定性测试仪"],
}


def _match_equipment(row_label: str, equipment: list[dict[str, Any]]) -> dict[str, Any] | None:
    target = _norm(row_label)
    if not target:
        return None
    candidates = [target]
    for key, synonyms in EQUIPMENT_SYNONYMS.items():
        if key in target or target in key:
            candidates.extend(_norm(item) for item in synonyms)
    for item in equipment:
        name = _norm(item.get("equipment_name", "") or item.get("设备名称", ""))
        if any(candidate and (candidate in name or name in candidate) for candidate in candidates):
            return item
    return None


def _equipment_raw_value(field: dict[str, Any], item: dict[str, Any], operator: str) -> str:
    header = _norm(field.get("col_header", "") or field.get("label", ""))
    if "型号规格" in header or "技术要求" in header:
        return str(item.get("model", "") or item.get("型号规格", ""))
    if "准确度" in header or "测量不确定度" in header or "测量范围" in header:
        return str(item.get("measuring_range", "") or item.get("测量范围", ""))
    if "管理编号" in header or "设备编号" in header:
        return str(item.get("management_no", "") or item.get("管理编号", ""))
    if "校准核查证书编号" in header:
        return str(item.get("calibration_certificate", "") or item.get("校准证书编号", ""))
    if "溯源机构" in header:
        return str(item.get("traceability_agency", "") or item.get("溯源机构", ""))
    if "有效期至" in header or "校准核查有效期" in header or "标准块有效期" in header:
        return str(item.get("calibration_due", "") or item.get("校准有效期至", ""))
    if "使用前状态" in header or "状态确认" in header:
        status = str(item.get("usage_status", "正常") or "正常")
        return "☑正常 □异常" if status == "正常" else "□正常 ☑异常"
    if "确认人" in header or "核查人" in header:
        return operator
    if "设备器具名称" in header or header == "名称":
        return str(item.get("equipment_name", "") or item.get("设备名称", ""))
    return ""


def _auto_checkbox(original: str, combined: str, context: dict[str, Any]) -> str:
    if "检验类别" in combined:
        return _select_checkbox(original, "委托检测")
    if "接收状态" in combined or "样品状态" in combined:
        return _select_checkbox(original, "完好")
    if "检测依据" in combined or "检测方法" in combined:
        method = str(context.get("method_code", "") or context.get("standard", ""))
        for token in re.findall(r"[A-Za-z]+(?:/[A-Za-z]+)?\s*\d+(?:\.\d+)?(?:-\d+)?", method):
            candidate = _select_checkbox(original, token.strip())
            if candidate != original:
                return candidate
    if "材料工艺" in combined or "材料名称" in combined:
        material = str(context.get("material", ""))
        candidate = _select_checkbox(original, material)
        return candidate if candidate != original else _select_other(original, material)
    if "样品规格" in combined or "型号规格" in combined:
        model = str(context.get("model", ""))
        candidate = _select_checkbox(original, model)
        return candidate if candidate != original else _select_other(original, model)
    return original


def prefill_template_fields(
    template_name: str,
    context: dict[str, Any],
    equipment: list[dict[str, Any]] | None = None,
    prior: dict[str, Any] | None = None,
) -> dict[str, str]:
    manifest = template_manifest(template_name)
    equipment = equipment or []
    prior = prior or {}
    values: dict[str, str] = {}
    samples = [str(x) for x in context.get("sample_nos", [])]
    sample_field_counter: dict[tuple[int, int], int] = {}
    sample_rows: dict[tuple[int, int], int] = {}

    for field in manifest:
        header = _norm(field.get("col_header", ""))
        if header in {"样品编号", "试样编号"} and field["row"] > 0:
            key = (field["table"], field["col"])
            index = sample_field_counter.get(key, 0)
            sample_field_counter[key] = index + 1
            sample_rows[(field["table"], field["row"])] = index

    for field in manifest:
        key = field["key"]
        if key in prior:
            values[key] = str(prior.get(key, ""))
            continue
        original = str(field.get("template_text", "") or "")
        label_norm = _norm(field.get("label", ""))
        row_norm = _norm(field.get("row_label", ""))
        header_norm = _norm(field.get("col_header", ""))
        section_norm = _norm(field.get("section", ""))
        combined = label_norm + row_norm + header_norm
        raw = ""

        sample_row_index = sample_rows.get((field["table"], field["row"]))
        if sample_row_index is not None:
            if sample_row_index >= len(samples):
                values[key] = "不适用"
                continue
            if header_norm in {"样品编号", "试样编号"}:
                raw = samples[sample_row_index]
            elif "检测人员" in header_norm or "记录人" in header_norm:
                raw = context.get("operator", "")
            elif "材料名称" in header_norm:
                raw = context.get("material", "")

        if not raw:
            equipment_item = _match_equipment(field.get("row_label", ""), equipment)
            if equipment_item:
                raw = _equipment_raw_value(field, equipment_item, str(context.get("operator", "")))

        if not raw:
            if "委托单位地址" in combined:
                raw = context.get("client_address", "")
            elif "委托单位" in combined:
                raw = context.get("client_name", "")
            elif "生产单位" in combined or "生产厂家" in combined:
                raw = context.get("production_unit", "")
            elif "报告编号" in combined:
                raw = context.get("report_no", "")
            elif "产品编号批号" in combined or "样品编号批号" in combined:
                raw = context.get("product_no", "")
            elif "实验室样品编号" in combined:
                raw = "、".join(samples)
            elif "样品数量" in combined or "试样数量" in combined:
                raw = str(context.get("sample_quantity", len(samples)))
            elif "样品名称" in combined:
                raw = context.get("sample_name", "")
            elif "样品规格型号" in combined or label_norm == "样品规格" or "型号规格" in combined:
                raw = context.get("model", "")
            elif "材料工艺" in combined or label_norm == "材料名称":
                raw = context.get("material", "")
            elif "接收日期" in combined:
                raw = context.get("received_date", "")
            elif "检测日期" in combined or "测量日期" in combined or label_norm == "日期":
                raw = context.get("test_date", "")
            elif "检测地点" in combined:
                raw = context.get("detection_location", "")
            elif "检测依据" in combined and "□" not in original:
                raw = context.get("standard", "")
            elif "实验任务编号" in combined or "任务编号" in combined:
                raw = context.get("task_no", "")
            elif "检测人员" in combined or "记录人" in combined or "操作人" in combined:
                raw = context.get("operator", "")
            elif "核验人员" in combined or "复核人员" in combined:
                raw = context.get("reviewer", "")
            elif ("异常" in section_norm or "偏离" in section_norm) and not original:
                raw = "无"
            elif ("附件归档" in section_norm or "数据归档" in section_norm) and not original:
                raw = "详见内部实验数据追溯Excel"

        if "□" in original or "☐" in original:
            auto = _auto_checkbox(original, combined, context)
            values[key] = auto
        else:
            values[key] = _compose_cell_text(original, raw)
    return values


def _checkbox_complete(original: str, value: str) -> bool:
    if "☑" not in value:
        return False
    # Blank text attached to an unselected alternative is allowed. If an
    # "other/abnormal" option is selected, its text must also be completed.
    if re.search(r"☑\s*(其他|异常|有)\s*[：:]?\s*(_{2,}|＿{2,}|…{2,})", value):
        return False
    return True


def validate_template_fields(template_name: str, values: dict[str, Any]) -> list[dict[str, str]]:
    missing: list[dict[str, str]] = []
    for field in template_manifest(template_name):
        value = _clean(values.get(field["key"], ""))
        original = str(field.get("template_text", "") or "")
        complete = bool(value)
        if complete and ("待填写" in value or "待确认" in value):
            complete = False
        if complete and ("□" in original or "☐" in original):
            complete = _checkbox_complete(original, value)
        elif complete and BLANK_RE.search(value) and value not in {"无", "/", "不适用"}:
            complete = False
        if not complete:
            missing.append({
                "key": field["key"],
                "section": field["section"],
                "label": field["label"],
                "position": field["position"],
            })
    return missing


def completion_summary(template_name: str, values: dict[str, Any]) -> dict[str, int]:
    total = len(template_manifest(template_name))
    missing = len(validate_template_fields(template_name, values))
    return {"total": total, "completed": total - missing, "missing": missing}


def complete_demo_values(template_name: str, values: dict[str, Any]) -> dict[str, str]:
    """Populate short validation values; used only by package QA."""
    output = dict(values)
    for field in template_manifest(template_name):
        key = field["key"]
        original = str(field.get("template_text", "") or "")
        value = str(output.get(key, "") or "")
        if "□" in original or "☐" in original:
            if "☑" not in value:
                value = _select_checkbox(original, re.sub(r"^□\s*", "", original.split("□", 1)[1]).split("□", 1)[0].strip())
                if "☑" not in value:
                    value = original.replace("□", "☑", 1)
        else:
            if not value:
                value = "1"
            if BLANK_RE.search(value):
                value = BLANK_RE.sub("1", value)
        output[key] = value
    return output


def _clone_rpr(source_run, target_run) -> None:
    if source_run is not None and source_run._r.rPr is not None:
        target_run._r.insert(0, deepcopy(source_run._r.rPr))


def _clear_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def _write_cell_text(cell, original: str, value: Any, changed: bool = False) -> None:
    text = "" if value is None else str(value)
    paragraphs = list(cell.paragraphs)
    paragraph = paragraphs[0] if paragraphs else cell.add_paragraph()
    source_run = next((r for p in paragraphs for r in p.runs), None)
    for p in paragraphs:
        _clear_runs(p)
    if not changed:
        run = paragraph.add_run(text)
        _clone_rpr(source_run, run)
        run.font.color.rgb = BLACK
        return

    matcher = difflib.SequenceMatcher(a=original, b=text)
    for tag, a1, a2, b1, b2 in matcher.get_opcodes():
        segment = text[b1:b2]
        if not segment:
            continue
        run = paragraph.add_run(segment)
        _clone_rpr(source_run, run)
        run.font.color.rgb = BLACK if tag == "equal" else RED


def fill_exact_template(
    template_name: str,
    values: dict[str, Any],
    changed_keys: set[str] | None = None,
):
    path = TEMPLATE_DIR / template_name
    if not path.exists():
        raise FileNotFoundError(f"受控原始记录模板不存在：{template_name}")
    changed_keys = changed_keys or set()
    doc = Document(path)
    manifest_map = {field["key"]: field for field in template_manifest(template_name)}
    for key, value in values.items():
        field = manifest_map.get(key)
        if not field:
            continue
        table = doc.tables[field["table"]]
        row = table.rows[field["row"]]
        if field["col"] >= len(row.cells):
            raise IndexError(f"模板字段位置失效：{template_name}/{field['position']}")
        cell = row.cells[field["col"]]
        _write_cell_text(cell, str(field.get("template_text", "") or ""), value, key in changed_keys)
    # Do not add, delete, merge, split or resize any template structures.
    return doc
