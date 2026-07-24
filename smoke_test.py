# -*- coding: utf-8 -*-
from pathlib import Path
import tempfile, zipfile, re
from docx import Document

import lims_db
from constants import EXPERIMENTS
from experiment_engine import schema, initial_parameters, initial_rows, calculate_rows
from business_record_engine import (
    initialize_business_record, calculate_business_record, validate_business_record,
    business_to_template_fields, fixed_and_manual_fields, visible_row_fields,
)
from record_word_engine import export_record
from trace_excel_engine import build_internal_trace_workbook


def _structure(path):
    doc=Document(path)
    return {
        "paragraphs":len(doc.paragraphs),"tables":len(doc.tables),"sections":len(doc.sections),
        "table_rows":[len(t.rows) for t in doc.tables],"table_cols":[len(t.columns) for t in doc.tables],
    }


def _context(experiment, sample_ids):
    return {
        "client_name":"测试委托客户","client_address":"测试地址","production_unit":"测试生产单位",
        "product_no":"TEST-BATCH-001","sample_name":"测试试样","model":"25 mm×2 mm×2 mm",
        "material":"钴铬合金","sample_nos":sample_ids,"sample_quantity":len(sample_ids),
        "received_date":"2026-07-23","report_no":"WT20260723001","task_no":"BP20260723001-P01-T01",
        "test_date":"2026-07-23","detection_location":"性能检测室","standard":EXPERIMENTS[experiment]["std"],
        "method_code":EXPERIMENTS[experiment]["method"],"operator":"测试实验员","reviewer":"测试复核员",
    }


def _demo_business(kind, sample_ids):
    record=initialize_business_record(kind,sample_ids,"性能检测室",{})
    params=record["parameters"]
    params["start_time"]="2026-07-23T10:00:00"
    params["end_time"]="2026-07-23T10:30:00"
    # Fill every visible manual parameter with a realistic, non-empty QA value.
    for field in fixed_and_manual_fields(kind)[1]:
        key=field["key"];typ=field.get("type","text")
        if params.get(key) not in (None,""):
            continue
        if typ=="number":params[key]=1.0
        elif typ=="select":params[key]=(field.get("options") or ["正常"])[0]
        elif key in {"start_time","end_time"}:params[key]="2026-07-23 10:00:00"
        else:params[key]="已确认"
    rows=record["rows"]
    for index,row in enumerate(rows,1):
        for key,label,typ in visible_row_fields(kind):
            if typ=="calc":continue
            if row.get(key) not in (None,"") and not (isinstance(row.get(key),float) and row.get(key)==0.0):
                continue
            if typ=="number":row[key]=float(index+1)
            elif typ.startswith("select:"):row[key]=typ.split(":",1)[1].split("|")[0]
            elif key=="note":row[key]=""
            else:row[key]=f"记录{index}"
    record["parameters"]=params
    record["rows"]=calculate_rows(kind,rows)
    record["equipment_checks"]=[
        {"management_no":"BPGL-A021","equipment_name":"电子万能试验机","status":"正常","note":"","required":True},
        {"management_no":"BPGL-A035","equipment_name":"数显维氏硬度计","status":"正常","note":"","required":True},
        {"management_no":"BPGL-A036","equipment_name":"粗糙度仪","status":"正常","note":"","required":True},
    ]
    record["overall_status"]="正常完成";record["deviation"]="无";record["retest"]="否"
    record=calculate_business_record(kind,record)
    if not record.get("report_summary"):record["report_summary"]="尚未形成有效检验结果。"
    if not record.get("report_conclusion"):record["report_conclusion"]="符合"
    return record


def main():
    root=Path(__file__).parent
    assert len(EXPERIMENTS)==10
    assert "模板原内容" not in (root/"app.py").read_text(encoding="utf-8")
    exp_block=(root/"app.py").read_text(encoding="utf-8").split('elif page=="实验记录":',1)[1].split('elif page=="原始记录复核":',1)[0]
    assert "data_editor" not in exp_block
    assert "template_manifest" not in exp_block
    assert "下载内部实验数据追溯Excel" in (root/"app.py").read_text(encoding="utf-8")

    equipment=[
        {"management_no":"BPGL-A021","equipment_name":"电子万能试验机","model":"STS50K","measuring_range":"（0～50）kN；（0～2000）N","calibration_time":"2026.7","lifecycle_status":"启用","binding_role":"主设备","required":1},
        {"management_no":"BPGL-A035","equipment_name":"数显维氏硬度计","model":"HV-30Z","measuring_range":"HV10","calibration_time":"2026.7","lifecycle_status":"启用","binding_role":"主设备","required":1},
        {"management_no":"BPGL-A036","equipment_name":"粗糙度仪","model":"TR200","measuring_range":"±160 μm","calibration_time":"2026.7","lifecycle_status":"启用","binding_role":"主设备","required":1},
    ]
    attachments=[{"attachment_id":"ATT-DEMO-001","sample_no":"BP20260723001-01"}]
    sample_ids=["BP20260723001-01","BP20260723001-02","BP20260723001-03"]

    with tempfile.TemporaryDirectory() as td_raw:
        td=Path(td_raw)
        for experiment,cfg in EXPERIMENTS.items():
            kind=cfg["kind"];template=cfg["template"];template_path=root/"templates"/template
            before=_structure(template_path)
            business=_demo_business(kind,sample_ids)
            assert not validate_business_record(kind,business,equipment),(experiment,validate_business_record(kind,business,equipment))
            values=business_to_template_fields(template,kind,_context(experiment,sample_ids),equipment,business,attachments,{})
            assert any("测试委托客户" in str(v) for v in values.values()), experiment
            assert any("BP20260723001-01" in str(v) for v in values.values()), experiment
            record={"record_no":"DEMO-"+cfg["key"],"task_no":"DEMO-"+cfg["key"],"version":1,"experiment":experiment,"payload":{"template_fields":values,"business_record":business}}
            generated=td/f"{cfg['key']}_{template}"
            generated.write_bytes(export_record(record,template,[]).getvalue())
            assert generated.stat().st_size>5000
            assert before==_structure(generated),(template,before,_structure(generated))

        # Full workflow using Vickers hardness and the concise business record.
        lims_db.DB_PATH=td/"test.db";lims_db.ATTACHMENT_DIR=td/"attachments";lims_db.SIGNATURE_DIR=td/"signatures";lims_db.init_db()
        assert lims_db.authenticate("admin","admin123")
        assert len(lims_db.list_equipment(True))==88
        assert len(lims_db.current_config_overview())==10
        lims_db.add_organization({"org_code":"C001","org_name":"测试客户","short_name":"客户","is_client":True,"is_manufacturer":False,"is_contract_manufacturer":False,"address":"测试地址","contact":"联系人","phone":"13800000000","credit_code":"","notes":""},"admin")
        lims_db.add_organization({"org_code":"M001","org_name":"测试生产单位","short_name":"生产单位","is_client":False,"is_manufacturer":True,"is_contract_manufacturer":False,"address":"","contact":"","phone":"","credit_code":"","notes":""},"admin")
        orgs=lims_db.list_organizations();client=next(x for x in orgs if x["org_code"]=="C001");producer=next(x for x in orgs if x["org_code"]=="M001")
        method=lims_db.experiment_method_by_name("维氏硬度试验")
        lims_db.add_catalog({"sample_code":"S001","sample_name":"测试金属试样","model":"10 mm×10 mm×1 mm","material_name":"钴铬合金","category":"金属","unit":"件","experiment_codes":[method["experiment_code"]],"notes":""},"admin")
        catalog=next(x for x in lims_db.list_catalog() if x["sample_code"]=="S001");cn="WT20260723001"
        lims_db.create_commission({"commission_no":cn,"client_org_id":client["id"],"client_name":client["org_name"],"client_address":client["address"],"contact":client["contact"],"phone":client["phone"],"production_org_id":producer["id"],"production_org_name":producer["org_name"],"production_relation":"生产单位","commission_date":"2026-07-23","due_date":"2026-08-23","subcontract_allowed":"否","report_medium":"电子档","conformity_judgment":"是","uncertainty":"否","delivery_method":"Email","cnas_mark":"否","capability":"完全满足","notes":""},[{"group_no":"BP20260723001","catalog_id":catalog["id"],"sample_name":catalog["sample_name"],"model":catalog["model"],"material_name":catalog["material_name"],"product_no":"B001","quantity":3,"unit":"件","condition":"完好","condition_note":"","storage_area":"A区域","notes":"","experiment_codes":[method["experiment_code"]]}],"receiver")
        group=lims_db.commission_groups(cn)[0];pn=lims_db.create_task_package(group["id"],[method["experiment_code"]],"tester","reviewer","receiver")
        task=lims_db.package_tasks(pn)[0];lims_db.accept_package(pn,"tester","样品已收到，确认完好",{task["task_no"]:"显微检测室"},"正常")
        task=lims_db.task(task["task_no"]);assert task["detection_location"]=="显微检测室"
        lims_db.mark_task_experiment_time(task["task_no"],"tester","开始")
        lims_db.mark_task_experiment_time(task["task_no"],"tester","结束")
        task=lims_db.task(task["task_no"]);assert task["experiment_started_at"] and task["experiment_ended_at"]
        snapshot=lims_db.task_config_snapshot(task["task_no"]);assert snapshot["record_template_file"]=="RECORD_R011_VICKERS.docx"
        business=_demo_business("hv",task["sample_nos_list"])
        context=_context("维氏硬度试验",task["sample_nos_list"]);context.update({"task_no":task["task_no"]})
        values=business_to_template_fields(snapshot["record_template_file"],"hv",context,snapshot["equipment"],business,[],{})
        payload={"common":{"record_no":task["task_no"],"task_no":task["task_no"],"commission_no":cn,"report_no":cn,"client":"测试客户","sample_name":"测试金属试样","sample_no":"、".join(task["sample_nos_list"]),"model":"10 mm×10 mm×1 mm","material":"钴铬合金","method_code":snapshot["method_code"],"standard":snapshot["standard"],"test_date":"2026-07-23","operator":"测试实验员","reviewer":"测试复核员"},"business_record":business,"template_name":snapshot["record_template_file"],"template_fields":values,"equipment_snapshot":business["equipment_checks"],"deviation":"无","retest":"否","report_summary":business["report_summary"],"report_conclusion":business["report_conclusion"],"configuration_snapshot":snapshot}
        lims_db.save_record(task["task_no"],1,payload,"tester","待复核",snapshot["record_template_version"] or "A/0",snapshot["sop_version"] or "A/0")
        lims_db.review_record(task["task_no"],1,"reviewer","通过","通过")
        rec=lims_db.record(task["task_no"],1);final_doc=td/"final_record.docx";final_doc.write_bytes(export_record(rec,snapshot["record_template_file"],[]).getvalue())
        assert _structure(final_doc)==_structure(root/"templates"/snapshot["record_template_file"])
        lims_db.submit_package_return(pn,"tester",[{"sample_no":sample_no,"condition":"完好","note":""} for sample_no in task["sample_nos_list"]])
        lims_db.confirm_package_return(pn,"receiver",[{"sample_no":sample_no,"location":"A区域"} for sample_no in task["sample_nos_list"]])
        assert lims_db.dashboard_counts()["reports"]==1
        generated_report=lims_db.report(cn);assert generated_report and generated_report["status"]=="待检测员确认"
        attachment_id=lims_db.save_attachment({"commission_no":cn,"package_no":pn,"task_no":task["task_no"],"sample_no":task["sample_nos_list"][0],"attachment_type":"实验过程照片","original_name":"test.jpg","captured_at":"2026-07-23 10:00:00","description":"过程照片","is_original":True},b"test-image-content","tester")
        attachment_rows=lims_db.list_attachments(task_no=task["task_no"]);assert attachment_rows[0]["attachment_id"]==attachment_id;assert "equipment_software" not in attachment_rows[0]
        dynamic=build_internal_trace_workbook(cn);xlsx=td/"trace.xlsx";xlsx.write_bytes(dynamic.getvalue())
        with zipfile.ZipFile(xlsx) as z:
            assert z.testzip() is None
            combined="".join(z.read(name).decode("utf-8","ignore") for name in z.namelist() if name.endswith(".xml"))
        assert "软件或设备" not in combined and "设备或软件" not in combined
        assert "附件编号" in combined and "实验总台账" in combined

    print("SMOKE TEST PASSED: concise form UI, logical defaults, exact controlled templates, attachment separation and Excel download")

if __name__=="__main__":main()
