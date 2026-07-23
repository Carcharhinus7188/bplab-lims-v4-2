from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
TARGET = ROOT / "templates" / "report" / "检验报告_非受控测试母版.docx"


def set_cell(cell, text, bold=False, size=9):
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = bold
            run.font.name = "宋体"
            run.font.size = Pt(size)


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    cover = doc.add_table(rows=13, cols=2)
    cover.style = "Table Grid"
    cover.cell(0, 0).merge(cover.cell(0, 1))
    cover.cell(1, 0).merge(cover.cell(1, 1))
    cover.cell(2, 0).merge(cover.cell(2, 1))
    labels = ["", "", "检验报告 / TEST REPORT", "报告编号", "委托编号", "委托单位", "生产单位", "样品名称", "规格型号", "批号/生产日期", "样品编号", "检验日期", "报告发布日期"]
    for row, label in enumerate(labels):
        set_cell(cover.cell(row, 0), label, bold=row in {0, 1, 2}, size=16 if row < 3 else 10)
    doc.add_section(WD_SECTION.NEW_PAGE)

    title = doc.add_paragraph("主要设备")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    equipment = doc.add_table(rows=15, cols=7)
    equipment.style = "Table Grid"
    for col, text in enumerate(["设备名称", "型号/规格", "管理编号", "量程/准确度", "证书编号", "溯源机构", "有效期至"]):
        set_cell(equipment.cell(0, col), text, bold=True)

    doc.add_section(WD_SECTION.NEW_PAGE)
    title = doc.add_paragraph("检测环境")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    environment = doc.add_table(rows=11, cols=5)
    environment.style = "Table Grid"
    for col, text in enumerate(["检验项目", "检测地点", "温度", "相对湿度", "其他环境/异常"]):
        set_cell(environment.cell(0, col), text, bold=True)

    doc.add_section(WD_SECTION.NEW_PAGE)
    title = doc.add_paragraph("检验结果")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    results = doc.add_table(rows=11, cols=6)
    results.style = "Table Grid"
    for col, text in enumerate(["样品编号", "检验项目", "检验依据", "标准要求", "检验结果", "单项结论"]):
        set_cell(results.cell(0, col), text, bold=True)

    doc.add_section(WD_SECTION.NEW_PAGE)
    final = doc.add_table(rows=7, cols=4)
    final.style = "Table Grid"
    final.cell(0, 1).merge(final.cell(0, 3))
    final.cell(1, 1).merge(final.cell(1, 3))
    final.cell(2, 1).merge(final.cell(2, 3))
    final.cell(6, 1).merge(final.cell(6, 3))
    set_cell(final.cell(0, 0), "样品情况说明", bold=True)
    set_cell(final.cell(1, 0), "检验结论", bold=True)
    set_cell(final.cell(2, 0), "总体判定", bold=True)
    set_cell(final.cell(3, 0), "检测员", bold=True)
    set_cell(final.cell(3, 2), "日期", bold=True)
    set_cell(final.cell(4, 0), "核验员", bold=True)
    set_cell(final.cell(4, 2), "日期", bold=True)
    set_cell(final.cell(5, 0), "批准人", bold=True)
    set_cell(final.cell(5, 2), "日期", bold=True)
    set_cell(final.cell(6, 0), "文件状态", bold=True)

    TARGET.parent.mkdir(parents=True, exist_ok=True)
    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()

