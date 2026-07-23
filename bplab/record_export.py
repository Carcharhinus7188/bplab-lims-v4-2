from __future__ import annotations

import hashlib
from copy import deepcopy
from datetime import date
from pathlib import Path
from typing import Any, Callable

from docx import Document
from docx.document import Document as DocumentType
from docx.table import _Cell

from .config import OUTPUT_DIR, RECORD_TEMPLATE_DIR


def _text(value: Any, digits: int = 3) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "是" if value else "否"
    if isinstance(value, float):
        return f"{value:.{digits}f}"
    return str(value)


def _date_cn(value: str | None) -> str:
    if not value:
        return ""
    try:
        y, m, d = value[:10].split("-")
        return f"{int(y)}年{int(m)}月{int(d)}日"
    except (ValueError, IndexError):
        return value


def _choice(options: list[str], selected: str | None) -> str:
    return " ".join(f"{'☑' if option == selected else '□'}{option}" for option in options)


def _multi(options: list[str], selected: list[str] | None) -> str:
    selected = selected or []
    return " ".join(f"{'☑' if option in selected else '□'}{option}" for option in options)


def _write_cell(cell: _Cell, value: Any) -> None:
    text = _text(value)
    runs = [run for paragraph in cell.paragraphs for run in paragraph.runs]
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run.text = ""
    else:
        cell.paragraphs[0].add_run(text)


def _cell(doc: DocumentType, table: int, row: int, col: int) -> _Cell:
    return doc.tables[table].rows[row].cells[col]


def _put(doc: DocumentType, table: int, row: int, col: int, value: Any) -> None:
    _write_cell(_cell(doc, table, row, col), value)


def _equipment(task: dict[str, Any], keyword: str) -> dict[str, Any]:
    for item in task.get("equipment", []):
        if keyword in item.get("name", ""):
            return item
    return {}


def _equipment_line(item: dict[str, Any]) -> str:
    if not item:
        return ""
    parts = [item.get("name", ""), item.get("model", ""), item.get("management_no", "")]
    return " / ".join(part for part in parts if part)


def _structure_signature(doc: DocumentType) -> tuple[Any, ...]:
    section_sig = tuple(
        (
            section.page_width,
            section.page_height,
            section.top_margin,
            section.bottom_margin,
            section.left_margin,
            section.right_margin,
        )
        for section in doc.sections
    )
    table_sig = tuple(tuple(len(row.cells) for row in table.rows) for table in doc.tables)
    return len(doc.paragraphs), len(doc.tables), table_sig, section_sig


def _basic_context(task: dict[str, Any]) -> dict[str, Any]:
    inherited = task["inherited"]
    data = task["data"]
    return {
        "inherited": inherited,
        "record_no": task["task_no"],
        "report_no": f"{inherited['commission_no']}-R",
        "date": (task.get("started_at") or date.today().isoformat())[:10],
        "client": inherited.get("client_name", ""),
        "production": inherited.get("production_unit", ""),
        "sample_name": inherited.get("sample_name", ""),
        "sample_ids": inherited.get("sample_ids", []),
        "batch": inherited.get("batch_no", ""),
        "material": inherited.get("material", ""),
        "spec": inherited.get("specification", ""),
        "quantity": inherited.get("quantity", 0),
        "method": inherited.get("method", ""),
        "basis": inherited.get("basis", ""),
        "location": inherited.get("location", task.get("location", "")),
        "tester": task.get("tester_name", ""),
        "reviewer": task.get("reviewer_name", ""),
        "data": data,
        "calc": task["calculations"],
    }


def _fill_bending(doc: DocumentType, task: dict[str, Any]) -> None:
    c = _basic_context(task)
    d, calc = c["data"], c["calc"]
    env, run, p = d.get("environment", {}), d.get("run", {}), d.get("parameters", {})
    _put(doc, 0, 0, 1, c["record_no"])
    _put(doc, 0, 0, 3, c["report_no"])
    _put(doc, 0, 0, 5, _date_cn(c["date"]))
    _put(doc, 0, 1, 1, c["client"])
    _put(doc, 0, 1, 3, c["sample_name"])
    _put(doc, 0, 1, 5, c["batch"] or "、".join(c["sample_ids"]))
    _put(doc, 0, 2, 1, c["material"])
    _put(doc, 0, 2, 3, d.get("print_process", "SLM"))
    _put(doc, 0, 2, 5, d.get("heat_treatment_record", "不适用"))
    _put(doc, 0, 3, 1, _choice(["长轴平行z轴", "长轴垂直z轴（x/y轴）"], d.get("print_direction", "长轴垂直z轴（x/y轴）")))
    _put(doc, 0, 3, 3, c["quantity"])
    main = _equipment(task, "万能试验机")
    _put(doc, 1, 1, 1, _equipment_line(main))
    _put(doc, 1, 1, 2, _choice(["有效", "失效"], "有效"))
    _put(doc, 1, 1, 3, _choice(["正常", "异常"], "正常" if d.get("prechecks", {}).get("equipment_ok", True) else "异常"))
    _put(doc, 1, 1, 4, c["tester"])
    _put(doc, 1, 2, 1, run.get("sensor_id", ""))
    _put(doc, 1, 2, 2, _choice(["有效", "失效"], "有效"))
    _put(doc, 1, 2, 3, _choice(["已选择", "未选择"], "已选择"))
    _put(doc, 1, 2, 4, c["tester"])
    _put(doc, 1, 4, 1, run.get("software_version", ""))
    _put(doc, 1, 4, 3, _choice(["正常", "异常"], "正常"))
    _put(doc, 1, 6, 1, f"温度{_text(env.get('temperature_before'),1)}℃；湿度{_text(env.get('humidity_before'),1)}%RH")
    _put(doc, 1, 6, 3, _choice(["符合", "不符合"], "符合"))
    _put(doc, 2, 3, 2, f"校准值{_text(run.get('sensor_check'))}N；系数{_text(d.get('sensor_factor', 1))}")
    _put(doc, 3, 0, 1, task["task_no"])
    _put(doc, 3, 0, 3, len(d.get("samples", [])))
    _put(doc, 3, 1, 3, _text(p.get("span"), 3))
    _put(doc, 3, 3, 5, run.get("data_path", ""))
    _put(doc, 3, 4, 1, run.get("fixture_id", ""))
    _put(doc, 3, 4, 5, _choice(["合格", "不合格"], "合格"))
    _put(doc, 3, 5, 1, _choice(["是", "否"], "是"))
    _put(doc, 3, 5, 3, _choice(["是", "否"], "是"))
    _put(doc, 3, 5, 5, _choice(["牢固", "松动"], "牢固"))
    for idx in range(1, 7):
        row = idx
        if idx <= len(calc.get("samples", [])):
            s = calc["samples"][idx - 1]
            values = [
                s.get("sample_id"), s.get("length"), s.get("width"), s.get("height"),
                s.get("span_actual"), s.get("speed_actual"), s.get("fmax"), s.get("offset_stress"),
                s.get("data_file"), _choice(["完整", "断裂", "异常"], "完整" if s.get("valid") == "有效" else "异常"),
                _choice(["符合", "不符合"], s.get("judgment") if s.get("judgment") in {"符合", "不符合"} else None),
                "",
            ]
        else:
            values = ["/"] * 12
        for col, value in enumerate(values, start=1):
            _put(doc, 4, row, col, value)
    _put(doc, 4, 7, 1, _choice(["全部符合", "存在不符合", "需复测/技术评审"], {
        "符合": "全部符合", "不符合": "存在不符合"
    }.get(calc.get("judgment"), "需复测/技术评审")))
    for row in range(1, 10):
        _put(doc, 5, row, 2, _choice(["是", "否"], "是"))
    _put(doc, 7, 0, 1, run.get("data_path", ""))
    _put(doc, 7, 0, 3, run.get("data_path", ""))
    _put(doc, 7, 3, 1, f"{c['tester']} / {_date_cn(c['date'])}")
    _put(doc, 7, 3, 3, f"{c['reviewer']} / {_date_cn(task.get('reviewed_at', '')[:10]) if task.get('reviewed_at') else ''}")
    _put(doc, 7, 4, 3, _choice(["合格", "不合格", "仅描述结果", "需复测"], {
        "符合": "合格", "不符合": "不合格", "仅记录": "仅描述结果"
    }.get(calc.get("judgment"), "需复测")))


def _fill_vickers(doc: DocumentType, task: dict[str, Any]) -> None:
    c = _basic_context(task)
    d, calc = c["data"], c["calc"]
    env, run, p = d.get("environment", {}), d.get("run", {}), d.get("parameters", {})
    main = _equipment(task, "维氏硬度计")
    block = _equipment(task, "标准维氏硬度块")
    _put(doc, 0, 0, 1, c["record_no"])
    _put(doc, 0, 0, 3, _date_cn(c["date"]))
    _put(doc, 0, 1, 1, c["method"])
    _put(doc, 0, 2, 1, c["client"])
    _put(doc, 0, 2, 3, c["sample_name"])
    _put(doc, 0, 2, 5, c["batch"] or "、".join(c["sample_ids"]))
    _put(doc, 0, 3, 1, f"{c['quantity']} 件")
    _put(doc, 0, 3, 3, c["location"])
    _put(doc, 0, 3, 5, c["report_no"])
    _put(doc, 1, 0, 1, _text(env.get("temperature_before"), 1))
    _put(doc, 1, 0, 3, _text(env.get("humidity_before"), 1))
    _put(doc, 1, 0, 5, _choice(["是", "否"], "是"))
    _put(doc, 1, 1, 1, main.get("model", ""))
    _put(doc, 1, 1, 3, main.get("management_no", ""))
    _put(doc, 1, 1, 5, _date_cn(main.get("valid_until", "")))
    _put(doc, 1, 2, 1, _choice(["正常", "异常"], "正常"))
    _put(doc, 1, 2, 3, _choice(["清晰", "异常"], "清晰"))
    _put(doc, 1, 2, 5, run.get("software_version", ""))
    _put(doc, 1, 3, 1, run.get("standard_block_id") or block.get("management_no", ""))
    _put(doc, 1, 3, 3, run.get("standard_block_nominal", block.get("model", "")))
    _put(doc, 1, 3, 5, _date_cn(block.get("valid_until", "")))
    _put(doc, 1, 4, 3, run.get("block_m1"))
    _put(doc, 1, 4, 5, run.get("block_m2"))
    _put(doc, 1, 5, 1, run.get("block_m3"))
    _put(doc, 1, 5, 3, calc.get("summary", {}).get("standard_check_average", ""))
    _put(doc, 1, 5, 5, _choice(["合格", "不合格"], "合格"))
    methods = run.get("surface_confirm_method") or ["目视/显微镜确认"]
    _put(doc, 1, 6, 3, _multi(["目视/显微镜确认", "粗糙度仪实测", "制样工艺确认", "其他"], methods))
    _put(doc, 1, 6, 5, f"Ra={_text(run.get('surface_ra'))} μm / 编号：{run.get('surface_record_id','')}")
    _put(doc, 1, 7, 1, _multi(["平整", "清洁", "无油污", "无氧化皮", "无影响压痕缺陷"], ["平整", "清洁", "无油污", "无氧化皮", "无影响压痕缺陷"]))
    _put(doc, 1, 7, 3, _choice(["否", "是"], "否"))
    _put(doc, 1, 7, 5, _multi(["牢固", "测试面与压头轴线垂直"], ["牢固", "测试面与压头轴线垂直"]))
    _put(doc, 1, 8, 1, run.get("data_path", ""))
    _put(doc, 2, 0, 1, p.get("scale", "HV10"))
    _put(doc, 2, 0, 3, p.get("force", 98.07))
    _put(doc, 2, 0, 5, p.get("hold_time", 15))
    _put(doc, 2, 0, 7, _choice(["正常", "异常"], "正常"))
    for row in range(1, 5):
        _put(doc, 3, row, 1, _choice(["已确认", "不符合"], "已确认") if row < 4 else _choice(["符合要求", "不符合"], "符合要求"))
    _put(doc, 3, 5, 1, f"确认人员：{c['tester']} 日期：{_date_cn(c['date'])}")
    samples = calc.get("samples", [])
    for idx in range(6):
        sample = samples[idx] if idx < len(samples) else None
        for face in (1, 2):
            row = 1 + idx * 2 + (face - 1)
            if sample:
                values = [
                    sample.get("sample_id"),
                    f"面{face}",
                    sample.get(f"face{face}_hv1"),
                    sample.get(f"face{face}_hv2"),
                    sample.get(f"face{face}_hv3"),
                    sample.get(f"face{face}_average"),
                    _choice(["符合", "实测Ra"], "符合" if sample.get(f"face{face}_surface") == "合格" else "实测Ra"),
                    p.get("force", 98.07),
                    p.get("hold_time", 15),
                    sample.get("report_id", ""),
                    "",
                ]
            else:
                values = ["/"] * 11
            for col, value in enumerate(values):
                _put(doc, 4, row, col, value)
    for idx in range(6):
        row = idx + 1
        if idx < len(samples):
            sample = samples[idx]
            values = [
                sample.get("sample_id"), sample.get("face1_average"), sample.get("face2_average"),
                sample.get("overall_hv10"), _choice(["产品标准", "委托协议", "注册技术要求", "仅描述"], "仅描述" if calc.get("judgment") == "仅记录" else "委托协议"),
                calc.get("standard_requirement"), _choice(["符合", "不符合", "不判定"], {
                    "符合": "符合", "不符合": "不符合", "仅记录": "不判定"
                }.get(sample.get("judgment"), "不判定")), sample.get("report_id", ""), "",
            ]
        else:
            values = ["/"] * 9
        for col, value in enumerate(values):
            _put(doc, 6, row, col, value)
    _put(doc, 9, 0, 1, c["tester"])
    _put(doc, 9, 0, 3, _date_cn(c["date"]))
    _put(doc, 9, 0, 5, c["reviewer"])
    _put(doc, 9, 1, 5, _choice(["有效", "无效", "需技术评审"], "有效" if calc.get("judgment") in {"符合", "不符合", "仅记录"} else "需技术评审"))


def _fill_warpage(doc: DocumentType, task: dict[str, Any]) -> None:
    c = _basic_context(task)
    d, calc = c["data"], c["calc"]
    env, run = d.get("environment", {}), d.get("run", {})
    img = _equipment(task, "二次元")
    cutter = _equipment(task, "切割机")
    _put(doc, 0, 0, 1, c["record_no"])
    _put(doc, 0, 1, 1, c["client"])
    _put(doc, 0, 1, 3, c["report_no"])
    _put(doc, 0, 2, 1, c["sample_name"])
    _put(doc, 0, 2, 3, c["batch"] or c["spec"])
    _put(doc, 0, 3, 3, c["method"])
    _put(doc, 1, 0, 1, "、".join(c["sample_ids"]))
    _put(doc, 1, 1, 1, c["quantity"])
    _put(doc, 1, 2, 1, _multi(["已完成打印及后处理", "表面无污染", "无裂纹", "无影响测量缺陷", "编号已核对"], ["已完成打印及后处理", "表面无污染", "无裂纹", "无影响测量缺陷", "编号已核对"]))
    _put(doc, 1, 3, 1, c["location"])
    _put(doc, 1, 4, 1, f"{_text(env.get('temperature_before'),1)} ℃")
    _put(doc, 1, 5, 1, f"{_text(env.get('humidity_before'),1)} %RH")
    _put(doc, 1, 6, 1, _date_cn(c["date"]))
    _put(doc, 2, 1, 2, _choice(["合格", "不合格"], "合格"))
    _put(doc, 2, 1, 3, f"型号/编号：{img.get('model','')}/{img.get('management_no','')} 校准有效期至：{_date_cn(img.get('valid_until',''))}")
    _put(doc, 2, 3, 2, _choice(["合格", "不合格"], "合格"))
    _put(doc, 2, 3, 3, f"型号/编号：{cutter.get('model','')}/{cutter.get('management_no','')}")
    _put(doc, 2, 5, 3, f"切割片规格/批号：{run.get('cutting_disc','')}")
    _put(doc, 2, 6, 3, f"路径：{run.get('data_path','')}")
    for row in range(1, 6):
        _put(doc, 3, row, 2, _choice(["是", "否"], "是"))
        _put(doc, 7, row, 2, _choice(["是", "否"], "是"))
    samples = calc.get("samples", [])
    for idx in range(10):
        s = samples[idx] if idx < len(samples) else None
        row = idx + 1
        if s:
            before = [s.get("sample_id"), s.get("image_before"), _choice(["是", "否"], "是"), _choice(["是", "否"], "是"), s.get("h1"), c["tester"]]
            cut = [s.get("sample_id"), f"{s.get('cut_start','')} / {s.get('cut_end','')}", _choice(["是", "否"], "是"), _choice(["合格", "不合格"], "合格"), _choice(["是", "否"], s.get("recut", "否")), ""]
            after = [s.get("sample_id"), s.get("image_after"), _choice(["是", "否"], "是"), _choice(["是", "否"], "是"), s.get("h2"), c["tester"]]
            result = [s.get("sample_id"), s.get("h1"), s.get("h2"), s.get("delta_h"), calc.get("standard_requirement"), _choice(["合格", "不合格", "仅记录"], {"符合": "合格", "不符合": "不合格", "仅记录": "仅记录"}.get(s.get("judgment"), "仅记录")), ""]
        else:
            before = cut = after = ["/"] * 6
            result = ["/"] * 7
        for col, value in enumerate(before, start=1):
            _put(doc, 4, row, col, value)
        for col, value in enumerate(cut, start=1):
            _put(doc, 6, row, col, value)
        for col, value in enumerate(after, start=1):
            _put(doc, 8, row, col, value)
        for col, value in enumerate(result, start=1):
            _put(doc, 9, row, col, value)
    _put(doc, 13, 0, 1, c["tester"])
    _put(doc, 13, 0, 3, _date_cn(c["date"]))
    _put(doc, 13, 1, 1, c["reviewer"])
    _put(doc, 13, 1, 3, _date_cn(task.get("reviewed_at", "")[:10]) if task.get("reviewed_at") else "")


def _fill_thermal_expansion(doc: DocumentType, task: dict[str, Any]) -> None:
    c = _basic_context(task)
    d, calc = c["data"], c["calc"]
    env, run, p = d.get("environment", {}), d.get("run", {}), d.get("parameters", {})
    main = _equipment(task, "热膨胀")
    _put(doc, 0, 0, 1, c["record_no"])
    _put(doc, 0, 0, 3, c["report_no"])
    _put(doc, 0, 2, 3, "热膨胀系数测试")
    _put(doc, 0, 3, 1, c["location"])
    _put(doc, 0, 3, 3, _date_cn(c["date"]))
    _put(doc, 1, 0, 1, c["client"])
    _put(doc, 1, 0, 3, c["inherited"]["commission_no"] if "inherited" in c else "")
    _put(doc, 1, 1, 1, c["sample_name"])
    _put(doc, 1, 1, 3, "、".join(c["sample_ids"]))
    _put(doc, 1, 2, 1, _choice(["钛合金", "钴铬合金", "其他"], c["material"] if c["material"] in {"钛合金", "钴铬合金"} else "其他"))
    _put(doc, 1, 2, 3, f"{c['quantity']}件")
    _put(doc, 1, 3, 1, c["spec"])
    _put(doc, 1, 3, 3, _choice(["完好", "异常（说明）"], "完好"))
    _put(doc, 1, 5, 3, calc.get("standard_requirement", ""))
    for row, key, unit in [(1, "temperature_before", "℃"), (2, "humidity_before", "%RH")]:
        _put(doc, 2, row, 2, f"{_text(env.get(key),1)} {unit}")
        _put(doc, 2, row, 3, _choice(["是", "否"], "是"))
    for row, selected in [(3, "无"), (4, "正常"), (5, "正常"), (6, "已确认"), (7, "正常"), (8, "满足")]:
        options = ["无", "有"] if row == 3 else (["满足", "不满足"] if row == 8 else ["正常", "异常"])
        if row == 6:
            _put(doc, 2, row, 2, "☑已确认")
        else:
            _put(doc, 2, row, 2, _choice(options, selected))
        _put(doc, 2, row, 3, _choice(["是", "否"], "是"))
    _put(doc, 3, 1, 1, main.get("model", ""))
    _put(doc, 3, 1, 2, main.get("management_no", ""))
    _put(doc, 3, 1, 3, main.get("calibration_certificate", ""))
    _put(doc, 3, 1, 4, _date_cn(main.get("valid_until", "")))
    _put(doc, 3, 1, 5, _choice(["合格", "停用"], "合格"))
    _put(doc, 4, 6, 4, f"设定值：{_text(p.get('terminal_temperature'),1)} ℃")
    _put(doc, 4, 10, 4, f"PV：{_text(run.get('pv_value'),1)}")
    _put(doc, 4, 11, 4, f"路径：{run.get('data_path','')}")
    for row in range(1, 13):
        _put(doc, 4, row, 3, _choice(["是", "否"], "是"))
    samples = calc.get("samples", [])
    for idx in range(6):
        s = samples[idx] if idx < len(samples) else None
        row = idx + 1
        if s:
            dims = [s.get("sample_id"), c["material"], s.get("l0"), s.get("width_diameter"), s.get("thickness"), _choice(["正确", "不适用"], "正确"), s.get("initial_pv"), _choice(["是", "否"], "是"), ""]
            run_row = [s.get("sample_id"), s.get("start_time", ""), s.get("end_time", ""), s.get("start_temp"), s.get("end_temp"), _choice(["正常", "异常"], "正常"), _choice(["是", "否"], "是"), s.get("curve_file"), s.get("curve_file"), _choice(["有效", "无效"], s.get("valid", "有效"))]
            result = [s.get("sample_id"), f"{_text(s.get('start_temp'),1)}～{_text(s.get('end_temp'),1)}", s.get("l0"), s.get("delta_t"), s.get("delta_l_um"), s.get("alpha"), calc.get("standard_requirement"), _choice(["合格", "不合格", "仅记录"], {"符合": "合格", "不符合": "不合格", "仅记录": "仅记录"}.get(s.get("judgment"), "仅记录")), ""]
        else:
            dims, run_row, result = ["/"] * 9, ["/"] * 10, ["/"] * 9
        for col, value in enumerate(dims, start=1):
            _put(doc, 5, row, col, value)
        for col, value in enumerate(run_row, start=1):
            _put(doc, 6, row, col, value)
        for col, value in enumerate(result, start=1):
            _put(doc, 8, row, col, value)
    summary = calc.get("summary", {})
    _put(doc, 11, 0, 1, _choice(["合格", "不合格", "仅提供实测值", "需复测"], {"符合": "合格", "不符合": "不合格", "仅记录": "仅提供实测值"}.get(calc.get("judgment"), "需复测")))
    _put(doc, 11, 1, 1, f"{_text(summary.get('average_alpha'))} ×10⁻⁶/K")
    _put(doc, 11, 2, 1, f"{_text(summary.get('maximum_displacement_um'))} μm")
    _put(doc, 11, 3, 1, calc.get("report_result", ""))
    _put(doc, 11, 4, 1, f"签名：{c['tester']} 日期：{_date_cn(c['date'])}")
    _put(doc, 11, 4, 3, f"签名：{c['reviewer']} 日期：{_date_cn(task.get('reviewed_at','')[:10]) if task.get('reviewed_at') else ''}")


def _fill_thermal_shock(doc: DocumentType, task: dict[str, Any]) -> None:
    c = _basic_context(task)
    d, calc = c["data"], c["calc"]
    env, run = d.get("environment", {}), d.get("run", {})
    _put(doc, 0, 0, 1, c["client"])
    _put(doc, 0, 0, 3, c["inherited"]["commission_no"] if "inherited" in c else "")
    _put(doc, 0, 1, 1, c["report_no"])
    _put(doc, 0, 1, 3, c["record_no"])
    _put(doc, 0, 2, 1, c["sample_name"])
    _put(doc, 0, 2, 3, "、".join(c["sample_ids"]))
    _put(doc, 0, 3, 1, c["spec"])
    _put(doc, 0, 3, 3, c["batch"])
    _put(doc, 0, 4, 1, f"{c['quantity']} 颗")
    _put(doc, 0, 5, 1, _choice(["完好", "异常"], "完好"))
    _put(doc, 0, 5, 3, _date_cn(c["date"]))
    for row, before, after in [
        (1, env.get("temperature_before"), env.get("temperature_after")),
        (2, env.get("humidity_before"), env.get("humidity_after")),
        (3, env.get("illumination"), env.get("illumination")),
    ]:
        _put(doc, 1, row, 2, before)
        _put(doc, 1, row, 3, after)
        _put(doc, 1, row, 4, _choice(["是", "否"], "是"))
        _put(doc, 1, row, 5, c["tester"])
    _put(doc, 3, 2, 2, f"设定：100℃；稳定读数：{_text(run.get('oven_stable_temp'),1)}℃")
    _put(doc, 3, 3, 2, f"时长：{_text(run.get('first_heat_actual'),1)}min")
    _put(doc, 3, 5, 2, f"稳定读数：{_text(run.get('ice_temp_actual',1.0),1)}℃")
    _put(doc, 3, 6, 2, f"实际：{_text(run.get('transfer_actual'),1)}s")
    _put(doc, 3, 7, 2, f"时长：{_text(run.get('ice_actual'),1)}min")
    _put(doc, 3, 8, 2, f"温度：100℃；时长：{_text(run.get('second_heat_actual'),1)}min")
    for row in range(1, 9):
        _put(doc, 3, row, 3, _choice(["是", "否"], "是"))
    _put(doc, 6, 1, 2, f"环境温度：{_text(env.get('temperature_after'),1)}℃；☑无直吹风 ☑无阳光直射")
    _put(doc, 6, 2, 2, f"读数：{_text(run.get('cool_surface_temp'),1)}℃；稳定时间：30s")
    _put(doc, 6, 4, 2, f"照度：{_text(env.get('illumination'),0)}lx；放大镜：☑10× □其他")
    for row in range(1, 6):
        _put(doc, 6, row, 3, _choice(["是", "否"], "是"))
    samples = calc.get("samples", [])
    for idx in range(28):
        s = samples[idx] if idx < len(samples) else None
        table = 7 if idx < 14 else 8
        row = idx + 1 if idx < 14 else idx - 13
        if s:
            vals = [
                s.get("sample_id"), _choice(["无", "有"], s.get("initial_abnormal", "无")),
                _choice(["无", "有"], s.get("crack", "无")), _choice(["无", "有"], s.get("chipping", "无")),
                _choice(["无", "有"], s.get("fracture", "无")), _choice(["合格", "不合格"], "合格" if s.get("judgment") == "符合" else "不合格"),
                s.get("defect_detail", ""),
            ]
        else:
            vals = ["/"] * 7
        for col, value in enumerate(vals, start=1):
            _put(doc, table, row, col, value)
    summary = calc.get("summary", {})
    _put(doc, 9, 0, 1, f"{summary.get('total',0)}颗")
    _put(doc, 9, 0, 3, f"{summary.get('valid',0)}颗")
    _put(doc, 9, 1, 1, f"{summary.get('crack',0)}颗")
    _put(doc, 9, 1, 3, f"{summary.get('chipping',0)}颗")
    _put(doc, 9, 2, 1, f"{summary.get('fracture',0)}颗")
    _put(doc, 9, 3, 3, _choice(["合格", "不合格"], "合格" if calc.get("judgment") == "符合" else "不合格"))
    _put(doc, 9, 4, 1, calc.get("report_result", ""))
    _put(doc, 11, 1, 1, c["tester"])
    _put(doc, 11, 1, 3, _date_cn(c["date"]))
    _put(doc, 11, 2, 1, c["reviewer"])
    _put(doc, 11, 2, 3, _date_cn(task.get("reviewed_at", "")[:10]) if task.get("reviewed_at") else "")


def _fill_xray(doc: DocumentType, task: dict[str, Any]) -> None:
    c = _basic_context(task)
    d, calc = c["data"], c["calc"]
    env, run, p = d.get("environment", {}), d.get("run", {}), d.get("parameters", {})
    _put(doc, 0, 0, 3, c["record_no"])
    _put(doc, 0, 2, 3, c["report_no"])
    _put(doc, 0, 2, 5, _date_cn(c["date"]))
    _put(doc, 1, 0, 1, c["client"])
    _put(doc, 1, 0, 3, c["inherited"]["commission_no"] if "inherited" in c else "")
    _put(doc, 1, 1, 1, c["sample_name"])
    _put(doc, 1, 1, 3, "、".join(c["sample_ids"]))
    _put(doc, 1, 2, 1, c["quantity"])
    _put(doc, 1, 3, 3, c["location"])
    _put(doc, 1, 4, 1, _choice(["完好", "异常"], "完好"))
    _put(doc, 1, 4, 3, _multi(["清洁", "有污渍", "有水分/粉尘", "其他"], ["清洁"]))
    _put(doc, 1, 5, 1, _choice(["委托要求", "产品技术要求", "内部质量要求", "其他"], "委托要求"))
    _put(doc, 1, 5, 3, c["method"])
    _put(doc, 2, 1, 1, f"{_text(env.get('temperature_before'),1)} ℃")
    _put(doc, 2, 1, 5, _choice(["符合", "不符合"], "符合"))
    _put(doc, 2, 2, 1, f"{_text(env.get('humidity_before'),1)} %RH")
    _put(doc, 2, 2, 5, _choice(["符合", "不符合"], "符合"))
    _put(doc, 2, 5, 1, _multi(["警示标识正常", "防护门/设施有效", "警示灯正常", "联锁正常", "急停正常", "区域无无关人员"], ["警示标识正常", "防护门/设施有效", "警示灯正常", "联锁正常", "急停正常", "区域无无关人员"]))
    _put(doc, 2, 5, 5, _choice(["允许曝光", "禁止曝光"], "允许曝光"))
    for row in range(1, 9):
        if row - 1 < len(task.get("equipment", [])):
            eq = task["equipment"][row - 1]
            _put(doc, 3, row, 2, eq.get("model", ""))
            _put(doc, 3, row, 3, eq.get("management_no", ""))
            _put(doc, 3, row, 4, _date_cn(eq.get("valid_until", "")) or "待补充")
        _put(doc, 3, row, 6, _choice(["合格", "不合格"], "合格"))
    _put(doc, 4, 1, 0, run.get("density_strip_id", ""))
    _put(doc, 4, 1, 1, run.get("density_nominal", ""))
    _put(doc, 4, 1, 2, run.get("density_m1", ""))
    _put(doc, 4, 1, 3, run.get("density_m2", ""))
    _put(doc, 4, 1, 4, run.get("density_m3", ""))
    values = [run.get(f"density_m{i}") for i in (1, 2, 3) if run.get(f"density_m{i}") not in (None, "")]
    _put(doc, 4, 1, 5, sum(float(v) for v in values) / len(values) if values else "")
    _put(doc, 4, 1, 6, run.get("density_tolerance", ""))
    _put(doc, 4, 1, 7, _choice(["合格", "不合格"], "合格"))
    _put(doc, 4, 1, 8, c["tester"])
    _put(doc, 5, 0, 3, f"{_text(p.get('voltage'),1)} kV")
    _put(doc, 5, 1, 1, f"{_text(p.get('current'),1)} mA")
    _put(doc, 5, 1, 3, f"{_text(p.get('exposure_ms'),1)} ms")
    _put(doc, 5, 2, 1, f"{_text(p.get('mas'),1)} mAs")
    for row in range(0, 6):
        _put(doc, 5, row, 5, _choice(["符合", "调整"] if row < 4 else ["符合", "不符合"], "符合"))
    samples = calc.get("samples", [])
    for idx in range(10):
        row = idx + 1
        s = samples[idx] if idx < len(samples) else None
        if s:
            vals = [s.get("sample_id"), c["sample_name"], _choice(["完好", "异常"], "完好"), s.get("image_id"), "☑咬合面朝下", _choice(["清晰", "不清"], "清晰"), _choice(["有效", "无效"], s.get("image_valid", "有效")), _choice(["否", "是"], "否")]
        else:
            vals = ["/"] * 8
        for col, value in enumerate(vals, start=1):
            _put(doc, 6, row, col, value)
    for idx in range(6):
        row = idx + 1
        s = samples[idx] if idx < len(samples) else None
        if s:
            vals = [
                s.get("sample_id"), s.get("image_id"),
                f"ROI-1：{_text(s.get('roi1_average'))}；ROI-2：{_text(s.get('roi2_average'))}；ROI-3：{_text(s.get('roi3_average'))}",
                s.get("thickness_estimate"), s.get("thickness_estimate"),
                _choice(["否", "是"], "是" if s.get("out_of_range") else "否"), s.get("abnormal_detail", ""),
            ]
        else:
            vals = ["/"] * 7
        for col, value in enumerate(vals):
            _put(doc, 9, row, col, value)
    first = samples[0] if samples else {}
    _put(doc, 11, 0, 1, _choice(["有效", "无效"], first.get("image_valid", "有效")))
    _put(doc, 11, 0, 3, first.get("thickness_estimate", ""))
    _put(doc, 11, 1, 1, _choice(["未见明显异常", "见异常"], "未见明显异常" if calc.get("judgment") == "符合" else "见异常"))
    _put(doc, 11, 1, 3, _choice(["合格", "不合格", "需复检", "超出方法适用范围"], {"符合": "合格", "不符合": "不合格"}.get(calc.get("judgment"), "需复检")))
    _put(doc, 11, 2, 1, calc.get("report_result", ""))
    _put(doc, 14, 1, 1, c["tester"])
    _put(doc, 14, 1, 3, _date_cn(c["date"]))
    _put(doc, 14, 2, 1, c["reviewer"])
    _put(doc, 14, 2, 3, _date_cn(task.get("reviewed_at", "")[:10]) if task.get("reviewed_at") else "")


def _fill_color(doc: DocumentType, task: dict[str, Any]) -> None:
    c = _basic_context(task)
    d, calc = c["data"], c["calc"]
    env, run = d.get("environment", {}), d.get("run", {})
    _put(doc, 0, 0, 1, c["record_no"])
    _put(doc, 0, 0, 3, c["report_no"])
    _put(doc, 0, 1, 1, c["client"])
    _put(doc, 0, 1, 3, _date_cn(c["date"]))
    _put(doc, 0, 3, 1, c["location"])
    _put(doc, 0, 3, 3, _choice(["委托检测", "型式检验", "方法验证", "质量控制", "其他"], "委托检测"))
    _put(doc, 1, 0, 1, c["sample_name"])
    _put(doc, 1, 0, 3, c["spec"])
    _put(doc, 1, 1, 1, "、".join(c["sample_ids"]))
    _put(doc, 1, 1, 3, c["batch"])
    _put(doc, 1, 2, 1, c["quantity"])
    _put(doc, 1, 2, 3, _choice(["完好", "异常，说明"], "完好"))
    _put(doc, 1, 5, 1, _choice(["产品标准", "委托协议", "注册技术要求", "仅描述不判定"], "委托协议"))
    _put(doc, 1, 5, 3, _choice(["无明显色泽差异为合格", "其他"], "无明显色泽差异为合格"))
    _put(doc, 2, 0, 1, f"{_text(env.get('temperature_before'),1)} ℃")
    _put(doc, 2, 0, 3, f"{_text(env.get('humidity_before'),1)} %RH")
    _put(doc, 2, 1, 1, _choice(["正常", "异常"], "正常"))
    _put(doc, 2, 1, 3, _multi(["清洁", "无彩色反光物", "无明显粉尘/振动"], ["清洁", "无彩色反光物", "无明显粉尘/振动"]))
    _put(doc, 2, 2, 1, _choice(["D65 灯箱", "等效光源"], "D65 灯箱"))
    _put(doc, 2, 2, 3, _multi(["清洁", "无污染", "无明显颜色反射"], ["清洁", "无污染", "无明显颜色反射"]))
    _put(doc, 2, 3, 3, _choice(["已确认", "未确认"], "已确认"))
    observers = [run.get(f"observer{i}", "") for i in (1, 2, 3)]
    for idx, name in enumerate(observers, start=1):
        _put(doc, 3, idx, 1, name)
        _put(doc, 3, idx, 3, _choice(["否", "是"], "否"))
        _put(doc, 3, idx, 4, _choice(["合格", "不合格"], "合格"))
        _put(doc, 3, idx, 5, name)
    _put(doc, 5, 0, 1, run.get("xenon_id", ""))
    _put(doc, 5, 0, 3, f"{_text(run.get('xenon_hours'),1)} h")
    _put(doc, 5, 1, 1, run.get("filter_id", ""))
    _put(doc, 5, 1, 3, f"{_text(run.get('filter_hours'),1)} h")
    _put(doc, 5, 2, 3, _choice(["是", "否"], "是"))
    lux_values = [run.get(f"surface_lux{i}") for i in (1, 2, 3)]
    _put(doc, 6, 1, 3, lux_values[0])
    _put(doc, 6, 1, 4, lux_values[1])
    _put(doc, 6, 1, 5, lux_values[2])
    valid_lux = [float(v) for v in lux_values if v not in (None, "")]
    _put(doc, 6, 1, 6, sum(valid_lux) / len(valid_lux) if valid_lux else "")
    _put(doc, 6, 1, 7, _choice(["是", "否"], "是"))
    _put(doc, 6, 1, 9, c["tester"])
    samples = calc.get("samples", [])
    for idx in range(12):
        s = samples[idx] if idx < len(samples) else None
        row = idx + 1
        if s:
            vals = [s.get("sample_id"), s.get("control_id", ""), _choice(["圆片", "牙形", "其他"], s.get("shape")), s.get("size"), _choice(["试样夹", "锡箔", "铝箔"], s.get("shield_method")), "一半", _choice(["是", "否"], "是"), s.get("position"), _choice(["是", "否"], "是"), ""]
        else:
            vals = ["/"] * 10
        for col, value in enumerate(vals, start=1):
            _put(doc, 8, row, col, value)
    observation_labels = ["未见明显差异", "轻微差异", "明显差异", "无法判定"]
    for idx, s in enumerate(samples[:6]):
        for observer_idx in range(3):
            row = 1 + idx * 3 + observer_idx
            result = s.get(f"observer{observer_idx + 1}_result")
            _put(doc, 12, row, 0, s.get("sample_id"))
            _put(doc, 12, row, 1, f"观察者{observer_idx + 1}")
            _put(doc, 12, row, 2, _choice(["合格", "不合格"], "合格"))
            _put(doc, 12, row, 3, s.get("comparison_note", ""))
            _put(doc, 12, row, 4, s.get("comparison_note", ""))
            _put(doc, 12, row, 5, _choice(observation_labels, result))
            _put(doc, 12, row, 6, observers[observer_idx])
            _put(doc, 12, row, 7, _date_cn(run.get("observation_date") or c["date"]))
    for idx in range(12):
        row = idx + 1
        s = samples[idx] if idx < len(samples) else None
        if s:
            vals = [
                s.get("sample_id"), s.get("observer1_result"), s.get("observer2_result"), s.get("observer3_result"),
                _choice(["未见明显色泽差异", "可见色泽差异", "无法判定"], "未见明显色泽差异" if s.get("majority_result") == "未见明显差异" else ("无法判定" if s.get("majority_result") == "无法判定" else "可见色泽差异")),
                _choice(["是", "否"], "是"), c["basis"],
                _choice(["合格", "不合格", "需复检", "仅描述"], {"符合": "合格", "不符合": "不合格", "无法判定": "需复检"}.get(s.get("judgment"), "仅描述")), "",
            ]
        else:
            vals = ["/"] * 9
        for col, value in enumerate(vals):
            _put(doc, 13, row, col, value)
    _put(doc, 14, 0, 1, _choice(["未见明显色泽差异", "可见轻微色泽差异", "可见明显色泽差异", "无法判定"], "未见明显色泽差异" if calc.get("judgment") == "符合" else "可见明显色泽差异"))
    _put(doc, 14, 0, 3, _choice(["合格", "不合格", "需复检", "不作符合性判定"], {"符合": "合格", "不符合": "不合格"}.get(calc.get("judgment"), "需复检")))
    _put(doc, 14, 1, 1, calc.get("report_result", ""))
    _put(doc, 17, 0, 1, c["tester"])
    _put(doc, 17, 0, 3, _date_cn(c["date"]))
    _put(doc, 17, 0, 5, c["reviewer"])


FILLERS: dict[str, Callable[[DocumentType, dict[str, Any]], None]] = {
    "bending": _fill_bending,
    "vickers": _fill_vickers,
    "warpage": _fill_warpage,
    "thermal_expansion": _fill_thermal_expansion,
    "thermal_shock": _fill_thermal_shock,
    "xray": _fill_xray,
    "color_stability": _fill_color,
}


def export_record(task: dict[str, Any], output_dir: Path | None = None) -> tuple[Path, dict[str, Any]]:
    config = task["config"]
    status = config.get("template_status")
    if status != "controlled":
        raise ValueError(config.get("template_note") or "当前实验缺少可用的受控DOCX母版")
    template = RECORD_TEMPLATE_DIR / config["record_template"]
    if not template.exists():
        raise FileNotFoundError(f"受控原始记录母版不存在：{template.name}")
    filler = FILLERS.get(task["experiment_code"])
    if not filler:
        raise ValueError(f"{config['name']}尚未配置精准Word映射")
    doc = Document(template)
    before = _structure_signature(doc)
    filler(doc, deepcopy(task))
    after_memory = _structure_signature(doc)
    if after_memory != before:
        raise RuntimeError("回填过程改变了母版结构，已阻止导出")
    target_dir = output_dir or OUTPUT_DIR
    target_dir.mkdir(parents=True, exist_ok=True)
    target = target_dir / f"{task['task_no']}_{config['name']}_原始记录.docx"
    doc.save(target)
    reopened = Document(target)
    after = _structure_signature(reopened)
    if after != before:
        target.unlink(missing_ok=True)
        raise RuntimeError("保存后母版结构校验失败，已删除异常输出")
    source_hash = hashlib.sha256(template.read_bytes()).hexdigest()
    output_hash = hashlib.sha256(target.read_bytes()).hexdigest()
    return target, {
        "template": template.name,
        "template_sha256": source_hash,
        "output_sha256": output_hash,
        "structure_preserved": True,
        "signature": repr(before),
    }
