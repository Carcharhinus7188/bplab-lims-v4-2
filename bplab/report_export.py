from __future__ import annotations

import hashlib
import json
from pathlib import Path
from typing import Any

from docx import Document

from .config import COMPANY_CN, COMPANY_EN, OUTPUT_DIR, REPORT_TEMPLATE_DIR
from .record_export import _date_cn, _put, _structure_signature


OFFICIAL_TEMPLATE = REPORT_TEMPLATE_DIR / "检验报告_受控母版.docx"
PREVIEW_TEMPLATE = REPORT_TEMPLATE_DIR / "检验报告_非受控测试母版.docx"


def official_template_available() -> bool:
    return OFFICIAL_TEMPLATE.exists()


def _period(start: str, end: str) -> str:
    if not start:
        return ""
    if not end or end == start:
        return _date_cn(start)
    return f"{_date_cn(start)}至{_date_cn(end)}"


def export_report(
    report: dict[str, Any],
    *,
    official: bool = False,
    output_dir: Path | None = None,
) -> tuple[Path, dict[str, Any]]:
    template = OFFICIAL_TEMPLATE if official else PREVIEW_TEMPLATE
    if official and not template.exists():
        raise ValueError("尚未提供受控检验报告Word母版，正式报告导出和批准已被阻止")
    if not template.exists():
        raise FileNotFoundError("非受控测试报告母版尚未生成")
    snapshot = report.get("snapshot")
    if snapshot is None:
        snapshot = json.loads(report.get("snapshot_json", "{}"))
    commission = snapshot.get("commission", {})
    groups = snapshot.get("groups", [])
    tasks = snapshot.get("tasks", [])
    doc = Document(template)
    before = _structure_signature(doc)

    _put(doc, 0, 0, 0, COMPANY_CN)
    _put(doc, 0, 1, 0, COMPANY_EN)
    _put(doc, 0, 3, 1, report["report_no"])
    _put(doc, 0, 4, 1, commission.get("commission_no", ""))
    _put(doc, 0, 5, 1, commission.get("client_name", ""))
    _put(doc, 0, 6, 1, commission.get("production_unit", ""))
    _put(doc, 0, 7, 1, "；".join(group.get("sample_name", "") for group in groups))
    _put(doc, 0, 8, 1, "；".join(group.get("specification", "") for group in groups))
    _put(doc, 0, 9, 1, "；".join(group.get("batch_no", "") for group in groups))
    _put(doc, 0, 10, 1, "；".join(
        "、".join(json.loads(group.get("sample_ids_json", "[]"))) for group in groups
    ))
    _put(doc, 0, 11, 1, _period(snapshot.get("test_date_start", ""), snapshot.get("test_date_end", "")))
    _put(doc, 0, 12, 1, report.get("published_at", "")[:10] if report.get("published_at") else "批准后生成")

    equipment = snapshot.get("equipment", [])
    for idx in range(14):
        row = idx + 1
        if idx < len(equipment):
            item = equipment[idx]
            values = [
                item.get("name", ""),
                item.get("model", ""),
                item.get("management_no", ""),
                item.get("measurement_range", ""),
                item.get("calibration_certificate", "") or "【基础数据待补充】",
                item.get("traceability_body", "") or "【基础数据待补充】",
                _date_cn(item.get("valid_until", "")) or "【基础数据待补充】",
            ]
        else:
            values = ["/"] * 7
        for col, value in enumerate(values):
            _put(doc, 1, row, col, value)

    environments = snapshot.get("environments", [])
    for idx in range(10):
        row = idx + 1
        if idx < len(environments):
            item = environments[idx]
            before_t, after_t = item.get("temperature_before"), item.get("temperature_after")
            before_h, after_h = item.get("humidity_before"), item.get("humidity_after")
            t = f"{before_t}～{after_t} ℃" if before_t != after_t else f"{before_t} ℃"
            h = f"{before_h}～{after_h} %RH" if before_h != after_h else f"{before_h} %RH"
            values = [item.get("experiment", ""), item.get("location", ""), t, h, item.get("other", "无异常")]
        else:
            values = ["/"] * 5
        for col, value in enumerate(values):
            _put(doc, 2, row, col, value)

    for idx in range(10):
        row = idx + 1
        if idx < len(tasks):
            task = tasks[idx]
            cfg = task.get("config", {})
            calc = task.get("calculations", {})
            inherited = task.get("inherited", {})
            values = [
                "、".join(inherited.get("sample_ids", [])),
                cfg.get("name", task.get("experiment_code", "")),
                cfg.get("method", ""),
                calc.get("standard_requirement", cfg.get("basis", "")),
                calc.get("report_result", ""),
                calc.get("judgment", ""),
            ]
        else:
            values = ["/"] * 6
        for col, value in enumerate(values):
            _put(doc, 3, row, col, value)

    _put(doc, 4, 0, 1, snapshot.get("sample_note", ""))
    _put(doc, 4, 1, 1, snapshot.get("final_conclusion", ""))
    _put(doc, 4, 2, 1, snapshot.get("final_judgment", ""))
    _put(doc, 4, 3, 1, commission.get("tester_name", "") if report.get("tester_signed_at") else "待检测员确认")
    _put(doc, 4, 3, 3, _date_cn(report.get("tester_signed_at", "")[:10]) if report.get("tester_signed_at") else "")
    _put(doc, 4, 4, 1, commission.get("reviewer_name", "") if report.get("reviewer_signed_at") else "待核验")
    _put(doc, 4, 4, 3, _date_cn(report.get("reviewer_signed_at", "")[:10]) if report.get("reviewer_signed_at") else "")
    _put(doc, 4, 5, 1, commission.get("approver_name", "") if report.get("approver_signed_at") else "待批准")
    _put(doc, 4, 5, 3, _date_cn(report.get("approver_signed_at", "")[:10]) if report.get("approver_signed_at") else "")
    _put(doc, 4, 6, 1, "正式受控报告" if official else "非受控测试预览：待替换实验室正式报告母版")

    if _structure_signature(doc) != before:
        raise RuntimeError("报告回填改变了母版结构，已阻止导出")
    target_dir = output_dir or OUTPUT_DIR
    target_dir.mkdir(parents=True, exist_ok=True)
    suffix = "正式检验报告" if official else "检验报告_非受控测试预览"
    target = target_dir / f"{report['report_no']}_{suffix}.docx"
    doc.save(target)
    if _structure_signature(Document(target)) != before:
        target.unlink(missing_ok=True)
        raise RuntimeError("报告保存后母版结构校验失败")
    return target, {
        "official": official,
        "template": template.name,
        "template_sha256": hashlib.sha256(template.read_bytes()).hexdigest(),
        "output_sha256": hashlib.sha256(target.read_bytes()).hexdigest(),
        "structure_preserved": True,
        "contains_see_original_record": "详见原始记录" in "\n".join(p.text for p in Document(target).paragraphs),
    }

